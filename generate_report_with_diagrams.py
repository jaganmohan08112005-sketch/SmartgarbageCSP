#!/usr/bin/env python3
"""
Generate the B.Tech Community Project Report for SmartGarbage.
Creates 4 matplotlib diagrams, embeds them into the DOCX at the correct locations.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os, sys

# ── Create diagrams first ──────────────────────────────────────────
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch

DIAG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_diagrams")
os.makedirs(DIAG_DIR, exist_ok=True)

C = {
    'dg': '#1B5E20', 'g': '#388E3C', 'lg': '#C8E6C9', 'pg': '#E8F5E9',
    'b': '#1565C0', 'lb': '#BBDEFB', 'pb': '#E3F2FD',
    'o': '#E65100', 'lo': '#FFE0B2', 'p': '#6A1B9A', 'lp': '#E1BEE7',
    'r': '#B71C1C', 'lr': '#FFCDD2', 'gr': '#424242', 'lgr': '#E0E0E0', 'w': '#FFFFFF',
}

def _box(ax, x, y, w, h, text, color, tc='white', fs=9, bold=False):
    p = FancyBboxPatch((x-w/2, y-h/2), w, h, boxstyle="round,pad=0.01,rounding_size=0.02",
                       facecolor=color, edgecolor=C['gr'], linewidth=1.2)
    ax.add_patch(p)
    ax.text(x, y, text, ha='center', va='center', fontsize=fs,
            fontweight='bold' if bold else 'normal', color=tc, multialignment='center')

def _arr(ax, x1, y1, x2, y2, color=None):
    ax.annotate('', xy=(x2,y2), xytext=(x1,y1),
                arrowprops=dict(arrowstyle='->', color=color or C['gr'], lw=1.5))

def _save(fig, name):
    p = os.path.join(DIAG_DIR, name)
    fig.savefig(p, dpi=200, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(fig)
    return p

def diag_architecture():
    fig, ax = plt.subplots(figsize=(10,7)); ax.set_xlim(0,10); ax.set_ylim(0,7.5); ax.axis('off')
    _box(ax,5,7.0,8,0.6,'USERS',C['dg'],fs=12,bold=True)
    for i,u in enumerate(['Citizen','Admin','Worker','Public']):
        _box(ax,1.5+i*2.3,7.0,1.6,0.4,u,C['g'],fs=8)
    _arr(ax,5,6.6,5,6.15)
    _box(ax,5,6.0,8,0.55,'BROWSER / MOBILE / PWA',C['b'],fs=10,bold=True)
    _arr(ax,5,5.65,5,5.2)
    _box(ax,5,5.05,5,0.4,'Service Worker Cache / CDN',C['lb'],tc=C['b'],fs=9,bold=True)
    _arr(ax,5,4.78,5,4.35)
    _box(ax,5,4.15,8.5,0.55,'FLASK APPLICATION (Gunicorn + Greenlet)',C['o'],fs=10,bold=True)
    _arr(ax,5,3.78,5,3.35)
    _box(ax,5,3.1,8.5,0.5,'',C['lo'])
    for i,m in enumerate(['Public\nPortal','Citizen\nPortal','Admin\nPortal','Worker\nPortal','IoT\nAPI']):
        _box(ax,0.85+i*1.8,3.1,1.5,0.45,m,C['o'],fs=7,bold=True)
    _arr(ax,5,2.75,5,2.3)
    _box(ax,5,2.1,8.5,0.5,'',C['lp'])
    for i,l in enumerate(['ML Engine','Job Queue\n(RQ)','Push\nNotif.','PAYT\nBilling']):
        _box(ax,1.3+i*2.3,2.1,1.8,0.45,l,C['p'],fs=7,bold=True)
    _arr(ax,5,1.75,5,1.35)
    _box(ax,5,1.15,8.5,0.5,'',C['pg'])
    for i,d in enumerate(['PostgreSQL\n(Supabase)','Redis','Sentry']):
        _box(ax,1.8+i*2.8,1.15,2.0,0.45,d,C['g'],fs=7,bold=True)
    _arr(ax,5,0.82,5,0.4)
    _box(ax,5,0.25,8.5,0.45,'EXTERNAL: Open-Meteo | Twilio | Razorpay | Render | GitHub',C['lgr'],tc=C['gr'],fs=8,bold=True)
    return _save(fig, 'architecture.png')

def diag_complaint():
    fig, ax = plt.subplots(figsize=(8,9)); ax.set_xlim(0,8); ax.set_ylim(0,9.5); ax.axis('off')
    _box(ax,4,9.0,4.5,0.55,'Citizen Reports Issue\n(GPS + Photo + Description)',C['g'],fs=9,bold=True)
    _arr(ax,4,8.65,4,8.2)
    _box(ax,4,8.0,3.5,0.5,'Server-side Validation',C['b'],fs=9,bold=True)
    _arr(ax,4,7.68,4,7.3)
    _box(ax,4,7.05,3.5,0.5,'Duplicate Detection',C['b'],fs=9,bold=True)
    _arr(ax,2.5,6.75,1.5,6.25)
    _arr(ax,5.5,6.75,6.5,6.25)
    _box(ax,1.5,6.0,2.0,0.5,'Duplicate →\nNotify',C['lr'],tc=C['r'],fs=8)
    _box(ax,6.5,6.0,2.2,0.5,'Valid →\nCreate Complaint',C['lg'],tc=C['g'],fs=8,bold=True)
    _arr(ax,6.5,5.65,6.5,5.2)
    _box(ax,6.5,5.0,2.8,0.5,'Admin Reviews &\nAssigns Worker',C['o'],fs=8,bold=True)
    _arr(ax,6.5,4.65,4,4.2)
    _box(ax,4,4.0,3.0,0.5,'Worker Dispatched\n(GPS Tracked)',C['p'],fs=9,bold=True)
    _arr(ax,4,3.65,4,3.2)
    _box(ax,4,3.0,3.5,0.5,'Worker Uploads\nPhoto + GPS Evidence',C['b'],fs=9,bold=True)
    _arr(ax,4,2.65,4,2.2)
    _box(ax,4,2.0,3.5,0.5,'Admin Verifies\nResolution',C['o'],fs=9,bold=True)
    _arr(ax,4,1.65,4,1.2)
    _box(ax,4,0.95,4.0,0.55,'✓ Resolved — Citizen Notified + Green Points',C['dg'],fs=9,bold=True)
    _box(ax,1.0,5.0,1.8,0.45,'Push/Email\nNotif.',C['lb'],tc=C['b'],fs=7)
    _arr(ax,1.9,5.0,2.5,4.0,color=C['b'])
    return _save(fig, 'complaint.png')

def diag_ml():
    fig, ax = plt.subplots(figsize=(10,6)); ax.set_xlim(0,10); ax.set_ylim(0,6.5); ax.axis('off')
    stages = [(1.0,4.5,'Available\nPrototype Data',C['lgr'],C['gr']),
              (2.8,4.5,'Data\nPreparation',C['b'],'white'),
              (4.6,4.5,'Feature\nEngineering',C['b'],'white'),
              (6.4,4.5,'Synthetic\nTraining Set\n(600 rows)',C['o'],'white'),
              (8.2,4.5,'RandomForest\nModel Training',C['p'],'white')]
    for x,y,t,c,tc in stages: _box(ax,x,y,1.6,1.0,t,c,tc=tc,fs=7,bold=True)
    for i in range(len(stages)-1): _arr(ax,stages[i][0]+0.8,4.5,stages[i+1][0]-0.8,4.5)
    stages2 = [(8.2,2.5,'Trained Model\n(Pickle)',C['p'],'white'),
               (6.0,2.5,'Prediction:\nHours Until\n90% Fill',C['g'],'white'),
               (3.8,2.5,'Rank Bins\nby Urgency',C['o'],'white'),
               (1.6,2.5,'Priority\nDispatch',C['dg'],'white')]
    for x,y,t,c,tc in stages2: _box(ax,x,y,1.6,1.0,t,c,tc=tc,fs=7,bold=True)
    for i in range(len(stages2)-1): _arr(ax,stages2[i][0]-0.8,2.5,stages2[i+1][0]+0.8,2.5)
    _arr(ax,8.2,3.95,8.2,3.1)
    _box(ax,3.0,1.0,5.5,0.55,'Fallback: Transparent heuristic when model unavailable',C['lb'],tc=C['b'],fs=8)
    _arr(ax,1.6,2.15,3.0,1.3,color=C['b'])
    ax.text(5,0.3,'Note: Synthetic data used during prototype development.\nReal-world historical data integration proposed for future work.',
            ha='center',va='center',fontsize=8,fontstyle='italic',color=C['gr'])
    return _save(fig, 'ml_pipeline.png')

def diag_pwa():
    fig, ax = plt.subplots(figsize=(10,6)); ax.set_xlim(0,10); ax.set_ylim(0,6.5); ax.axis('off')
    _box(ax,5,6.0,3.5,0.55,'User Submits Complaint',C['g'],fs=10,bold=True)
    _arr(ax,5,5.65,5,5.25)
    # Diamond
    diamond = plt.Polygon([(5,5.25),(6.3,4.9),(5,4.55),(3.7,4.9)],
                          facecolor=C['lo'],edgecolor=C['o'],linewidth=1.5)
    ax.add_patch(diamond)
    ax.text(5,4.9,'Internet\nAvailable?',ha='center',va='center',fontsize=8,fontweight='bold')
    # Yes
    _arr(ax,3.7,4.9,2.0,4.9,color=C['g'])
    ax.text(2.8,5.05,'Yes',ha='center',va='center',fontsize=8,color=C['g'],fontweight='bold')
    _box(ax,2.0,4.4,2.2,0.55,'Submit via API\n→ Database',C['g'],fs=8,bold=True)
    _arr(ax,2.0,4.05,2.0,3.55)
    _box(ax,2.0,3.3,2.2,0.5,'Success ✓\nNotified',C['lg'],tc=C['g'],fs=8)
    # No
    _arr(ax,6.3,4.9,8.0,4.9,color=C['r'])
    ax.text(7.2,5.05,'No',ha='center',va='center',fontsize=8,color=C['r'],fontweight='bold')
    _box(ax,8.0,4.4,2.2,0.55,'Store in\nIndexedDB',C['o'],fs=8,bold=True)
    _arr(ax,8.0,4.05,8.0,3.55)
    _box(ax,8.0,3.3,2.2,0.5,'Background\nSync',C['lo'],tc=C['o'],fs=8)
    _arr(ax,8.0,3.0,5.5,2.5)
    ax.text(6.8,2.75,'Connection\nRestored',ha='center',va='center',fontsize=7,color=C['gr'])
    _box(ax,5,2.2,3.0,0.55,'Service Worker\nSync Event',C['b'],fs=8,bold=True)
    _arr(ax,5,1.85,5,1.35)
    _box(ax,5,1.1,3.5,0.55,'API Submission\n→ Queue Emptied',C['g'],fs=8,bold=True)
    _arr(ax,5,0.75,5,0.3)
    _box(ax,5,0.1,3.0,0.4,'Citizen Notified ✓',C['dg'],fs=8,bold=True)
    return _save(fig, 'pwa_workflow.png')


# ── Generate diagrams ──────────────────────────────────────────────
print("Creating diagrams...")
DIAGS = {}
DIAGS['arch'] = diag_architecture()
DIAGS['complaint'] = diag_complaint()
DIAGS['ml'] = diag_ml()
DIAGS['pwa'] = diag_pwa()
print(f"Created {len(DIAGS)} diagrams")

# ── Document setup ──────────────────────────────────────────────────
doc = Document()
FONT_NAME = "Times New Roman"
section = doc.sections[0]
section.top_margin = Cm(2.54); section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54); section.right_margin = Cm(2.54)
section.page_height = Cm(29.7); section.page_width = Cm(21.0)

def _sf(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME; run.font.size = Pt(size)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rPr.insert(0, parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>'))
    else: rf.set(qn('w:eastAsia'), FONT_NAME)

def _p(text, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sa=6, sb=0, color=None):
    p = doc.add_paragraph(); p.alignment = align
    pf = p.paragraph_format; pf.space_after = Pt(sa); pf.space_before = Pt(sb); pf.line_spacing = 1.0
    run = p.add_run(text); _sf(run, size, bold, italic, color)
    return p

def _h1(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format; pf.space_before = Pt(24); pf.space_after = Pt(12); pf.line_spacing = 1.0
    run = p.add_run(text.upper()); _sf(run, 16, True); return p

def _h2(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format; pf.space_before = Pt(18); pf.space_after = Pt(8); pf.line_spacing = 1.0
    run = p.add_run(text); _sf(run, 14, True); return p

def _h3(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format; pf.space_before = Pt(12); pf.space_after = Pt(6); pf.line_spacing = 1.0
    run = p.add_run(text); _sf(run, 12, True); return p

def _body(text): return _p(text, 12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
def _bodyb(text): return _p(text, 12, bold=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
def _cap(text): return _p(text, 12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
def _pb(): doc.add_page_break()

def _img(path, width=Inches(5.5)):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.space_before = Pt(6); pf.space_after = Pt(6)
    run = p.add_run(); run.add_picture(path, width=width)
    return p

def _tbl(headers, rows, cw=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); _sf(run, 12, True)
        sh = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3"/>'); c._tc.get_or_add_tcPr().append(sh)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val)); _sf(run, 12)
    if cw:
        for i,w in enumerate(cw):
            for row in t.rows: row.cells[i].width = Inches(w)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════
for _ in range(3): doc.add_paragraph()
_p("COMMUNITY PROJECT REPORT", 16, True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=24)
_p("SMARTGARBAGE CHINTALAVALASA", 18, True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6, color=(0,100,0))
_p("Community-Based Smart Waste Management\nand Digital Governance System", 14, align=WD_ALIGN_PARAGRAPH.CENTER, sa=36)
_p("Submitted by", 12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
students = [
    ("MOPADA JAGANMOHAN", "2433144441"),
    ("LATCHUPATULA RESHMA", "24331A4434"),
    ("PATI NARASIMHA MURTHY", "2433144446"),
    ("KADA AUGUSTTN PAUL KUMAR", "24331A4426"),
]
for name,reg in students:
    _p(f"{name}  ({reg})", 12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
doc.add_paragraph()
_p("In partial fulfillment for the award of the degree of", 12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
_p("BACHELOR OF TECHNOLOGY", 14, True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)
_p("IN", 12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)
_p("COMPUTER SCIENCE & ENGINEERING", 14, True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
_p("(Data Science)", 12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=24)
_p("Under the esteemed Guidance of", 12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)
_p("Mrs. S. Nikhila", 14, True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
_p("Assistant Professor", 12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=24)
doc.add_paragraph()
_p("DEPARTMENT OF DATA ENGINEERING", 12, True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)
_p("MAHARAJ VIJAYARAM GAJAPATHI RAJ COLLEGE OF ENGINEERING (Autonomous)", 12, True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=3)
_p("(Approved by AICTE, New Delhi, and permanently affiliated to JNTUGV, Vizianagaram), Listed u/s 2(f) & 12(B) of UGC Act 1956.", 10, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)
_p("Vijayaram Nagar Campus, Chintalavalasa, Vizianagaram-535005, Andhra Pradesh", 10, align=WD_ALIGN_PARAGRAPH.CENTER, sa=6)
_p("October, 2025", 12, True, align=WD_ALIGN_PARAGRAPH.CENTER)

# CERTIFICATE
_pb(); _h1("CERTIFICATE"); doc.add_paragraph()
_body('This is to certify that the project entitled "SmartGarbage Chintalavalasa — Community-Based Smart Waste Management and Digital Governance System" is the bonafide work carried out by Mopada Jaganmohan (2433144441), Latchupatula Reshma (24331A4434), Pati Narasimha Murthy (2433144446), and Kada Augusttn Paul Kumar (24331A4426), of B.Tech V Sem CSE-DS, M.V.G.R. College of Engineering (Autonomous), Vizianagaram, during the year 2025-2026, in partial fulfilment of the requirements for the award of the Degree of Bachelor of Technology and that the project has not formed the basis for the award previously of any degree or any other similar title.')
doc.add_paragraph()
_p("Signature of Project Guide", 12, True, sa=3); _p("Mrs. S. Nikhila\nAssistant Professor\nDepartment: Data Engineering", 12, sa=24)
_p("Signature of Head of the Department", 12, True, sa=3); _p("Dr. Jyothi\nHead of the Department\nDepartment: Data Engineering", 12)

# DECLARATION
_pb(); _h1("DECLARATION"); doc.add_paragraph()
_body('We hereby declare that the work done on the dissertation entitled "SmartGarbage Chintalavalasa — Community-Based Smart Waste Management and Digital Governance System" has been carried out by us and submitted in partial fulfilment for the award of credits in Bachelor of Technology in Computer Science and Engineering (Data Science) of M.V.G.R College of Engineering (Autonomous) and affiliated to JNTUGV, Vizianagaram. The various contents incorporated in the dissertation have not been submitted for the award of any degree of any other institution or university.')
doc.add_paragraph()
for n in ["MOPADA JAGANMOHAN (2433144441)","LATCHUPATULA RESHMA (24331A4434)","PATI NARASIMHA MURTHY (2433144446)","KADA AUGUSTTN PAUL KUMAR (24331A4426)"]:
    _p(n, 12, sa=12)

# ACKNOWLEDGEMENT
_pb(); _h1("ACKNOWLEDGEMENT"); doc.add_paragraph()
_body("We express our sincere gratitude to our project guide for their invaluable guidance and support as our mentor throughout the project. Their unwavering commitment to excellence and constructive feedback motivated us to achieve our project goals. We are greatly indebted to them for their exceptional guidance.")
_body("Additionally, we extend our thanks to Prof. P.S. Sitharama Raju (Director), Dr. Y.M.C. Shekar (Principal), and Dr. Jyothi (Head of the Department) for their unwavering support and assistance, which were instrumental in the successful completion of the project.")
_body("We also acknowledge the dedicated assistance provided by all the staff members in the Department of Data Engineering. Finally, we appreciate the contributions of all those who directly or indirectly contributed to the successful execution of this endeavor.")

# ABSTRACT
_pb(); _h1("ABSTRACT"); doc.add_paragraph()
_body("Waste management in semi-urban Indian communities such as Chintalavalasa, Andhra Pradesh, faces significant challenges. Collection schedules are communicated informally, complaint tracking is paper-based, and citizens have no transparent mechanism to report issues or monitor resolution. Overflowing bins, delayed response times, and lack of accountability are common, with no data-driven approach to allocate resources or predict collection needs.")
_body("This project presents SmartGarbage, a community-based smart waste management system designed for the Chintalavalasa Gram Panchayat. The system provides a unified digital platform where citizens can check waste collection schedules, report missed pickups with photographic evidence and GPS location, and track complaint resolution in real time. Administrators can monitor operations through a live dashboard, assign workers, and view ward-level analytics. Sanitation workers receive dispatch notifications and can upload proof of completion directly from the field.")
_body("Beyond basic complaint management, the system introduces IoT-enabled smart bins that monitor fill levels, battery status, and environmental conditions, enabling proactive collection before overflow occurs. A machine learning module predicts bin overflow risk, allowing administrators to prioritize dispatch based on urgency. The platform also implements a pay-as-you-throw billing mechanism that charges residents based on the residual (non-segregated) waste they generate, incentivizing proper source segregation. A gamification feature called Green Points rewards citizens for consistent waste segregation behaviour.")
_body("The system is designed for low-connectivity environments through Progressive Web App capabilities that allow complaint filing and schedule access even without an internet connection, with automatic synchronization when connectivity is restored. Bilingual support in English and Telugu ensures accessibility across different literacy levels. Comprehensive security measures protect user data and system integrity.")
_body("The prototype has been deployed and tested for the Chintalavalasa community. While the current implementation uses synthetic data for machine learning training and simulated IoT telemetry — as real-world sensor hardware and historical data are not yet available at community scale — the architecture is designed for seamless integration with physical sensors and real usage data. The project demonstrates how low-cost, open-source digital infrastructure can bring transparency, accountability, and data-driven governance to waste management in semi-urban Indian panchayats.")
_p("Keywords: Smart Waste Management, Community Governance, IoT, Machine Learning, Digital India, PWA, PAYT Billing", 12, italic=True, sa=12)

# ABBREVIATIONS
_pb(); _h1("LIST OF ABBREVIATIONS"); doc.add_paragraph()
abbr = [("AI","Artificial Intelligence"),("API","Application Programming Interface"),("ARIA","Accessible Rich Internet Applications"),
    ("CDN","Content Delivery Network"),("CSP","Content Security Policy"),("CSS","Cascading Style Sheets"),
    ("DBMS","Database Management System"),("GPS","Global Positioning System"),("HTML","HyperText Markup Language"),
    ("HTTP","HyperText Transfer Protocol"),("IoT","Internet of Things"),("JSON-LD","JSON for Linked Data"),
    ("ML","Machine Learning"),("MFA","Multi-Factor Authentication"),("MVC","Model-View-Controller"),
    ("OTP","One-Time Password"),("OWASP","Open Web Application Security Project"),("PAYT","Pay-As-You-Throw"),
    ("PWA","Progressive Web App"),("RBAC","Role-Based Access Control"),("REST","Representational State Transfer"),
    ("RQ","Redis Queue"),("SBM","Swachh Bharat Mission"),("SEO","Search Engine Optimization"),
    ("SQL","Structured Query Language"),("SSL","Secure Sockets Layer"),("SW","Service Worker"),
    ("TTFB","Time to First Byte"),("VAPID","Voluntary Application Server Identification"),
    ("WCAG","Web Content Accessibility Guidelines"),("XML","Extensible Markup Language")]
_tbl(["Abbreviation","Full Form"], abbr, [1.5,5.0])

# ════════════════════════════════════════════════════════════════════
# CHAPTER 1
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("1. INTRODUCTION")
_h2("1.1 Problem Statement")
_body("Chintalavalasa is a semi-urban panchayat in Vizianagaram district, Andhra Pradesh, with a population of approximately 12,000 residents across five administrative wards. The current waste management system relies on manual collection schedules communicated informally, resulting in inconsistent service delivery. Citizens have no reliable mechanism to report missed collections, and complaint resolution lacks transparency and accountability.")
_body("Key problems include: (a) lack of centralized collection schedules; (b) absence of digital grievance redressal; (c) no real-time bin monitoring; (d) no data-driven resource allocation; (e) limited citizen engagement in segregation; and (f) no usage-based billing accountability.")

_h2("1.2 Project Objective")
_body("The primary objective is to design, develop, and deploy a community-based smart waste management and digital governance system for Chintalavalasa Gram Panchayat. Specific objectives include:")
for o in ["Web platform for schedules, complaint filing and tracking","Admin dashboard for real-time monitoring and worker dispatch",
    "IoT smart-bin telemetry for fill levels and environmental monitoring","ML module for overflow risk prediction",
    "PAYT billing with UPI/Razorpay integration","PWA capabilities with offline support",
    "Comprehensive security aligned with OWASP recommendations","Bilingual support (English and Telugu)"]:
    _p(f"\u2022  {o}", 12, sa=4)

_h2("1.3 Scope of the Project")
_body("\u2022  In Scope: Web application with four portals; IoT telemetry integration; ML overflow prediction; PAYT billing; Green Points gamification; PWA offline support; push notifications; bilingual support; and security architecture.")
_body("\u2022  Out of Scope: Native mobile apps; physical IoT hardware manufacturing; external municipal API integration; blockchain-based carbon credits; production WhatsApp/SMS (integrated but requiring API keys); and community-scale IoT deployment.")

# ════════════════════════════════════════════════════════════════════
# CHAPTER 2
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("2. LITERATURE SURVEY")
_body("A comprehensive review of existing literature and systems was conducted to identify gaps that SmartGarbage addresses.")
_h2("2.1 Existing Waste Management Approaches")
_body("Traditional waste management in Indian semi-urban areas relies on manual collection with fixed schedules and paper-based registers. The Swachh Bharat Mission (SBM) Grameen Phase II promotes source segregation and digital monitoring, but implementation remains inconsistent at the panchayat level.")
_h2("2.2 Digital Waste Management Systems")
_body("SBM Urban provides complaint registration for urban areas but lacks IoT integration and offline capabilities. GOV.UK sets the benchmark for government digital services with strong accessibility compliance. SmartGarbage draws design inspiration from these platforms while adding domain-specific waste management features.")
_h2("2.3 IoT-Based Smart Waste Management")
_body("Anagnostopoulos et al. (2015) demonstrated that IoT-based waste monitoring using ultrasonic sensors can reduce collection costs by 30-50%. Kumar and Sharma (2021) identified that most solutions focus on individual components rather than integrated platforms. SmartGarbage addresses this by integrating IoT with complaint management, citizen engagement, and ML prediction.")
_h2("2.4 Machine Learning for Waste Prediction")
_body("Afshin et al. (2021) applied gradient boosting to predict waste generation. Chen et al. (2020) demonstrated that Random Forest achieves comparable accuracy for short-term prediction. SmartGarbage implements a RandomForest regressor for bin overflow prediction with transparent fallback heuristics.")
_h2("2.5 Accessibility and Government Standards")
_body("WCAG 2.1 Level AA provides international accessibility standards. OWASP Top 10 identifies critical web security risks. SmartGarbage implements ARIA landmarks, skip-to-content links, keyboard navigation, and all nine OWASP-recommended security headers.")
_h2("2.6 Research / Implementation Gap")
_body("Existing solutions address individual aspects of waste management. There is no integrated, low-cost platform combining citizen grievance reporting, collection scheduling, IoT monitoring, ML prediction, PAYT billing, offline support, and comprehensive security — all designed for semi-urban Indian communities. SmartGarbage fills this gap.")
_h2("2.7 Proposed Contribution")
_body("SmartGarbage contributes an integrated, open-source platform combining citizen engagement, administrative oversight, worker coordination, IoT monitoring, and ML prediction with offline-first PWA capabilities, PAYT billing with gamification, and strong security and accessibility compliance — suitable for semi-urban Indian panchayats.")

# ════════════════════════════════════════════════════════════════════
# CHAPTER 3
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("3. DATA GATHERING / DATA USED")
_h2("3.1 Study Area / Community Profile")
_body("Chintalavalasa is a semi-urban panchayat in Vizianagaram district, Andhra Pradesh, serving approximately 12,000 residents across five wards. Coordinates range from latitude 18.0552 to 18.0751 and longitude 83.4005 to 83.4201.")
_h2("3.2 Data Collection Methods")
_body("\u2022  Community Interaction: Interviews with residents, ward members, and sanitation workers.\n\u2022  Field Observation: On-site observation of collection routes and complaint handling.\n\u2022  Administrative Data: Ward boundaries and population estimates from the Gram Panchayat.\n\u2022  System-Generated Data: Synthetic test data for ML training and IoT simulation.\n\u2022  Public Records: SBM Grameen guidelines, GOV.UK documentation, WCAG standards.")

_h2("3.3 Data Sources")
_tbl(["Data Source","Type","Description","Use"],[
    ("Community Surveys","Qualitative","Resident interviews","Requirements"),
    ("Field Observations","Qualitative","On-site observation","System design"),
    ("Administrative Records","Semi-structured","Ward boundaries","Study area"),
    ("Synthetic ML Data","Quantitative","600-row grid","ML training"),
    ("Synthetic IoT Data","Quantitative","Simulated sensors","IoT testing"),
    ("Test User Data","Quantitative","Registration records","Testing")], [1.5,1.2,2.2,1.5])

_h2("3.4 Ward Information")
_tbl(["Ward","Name","Latitude","Longitude"],[
    ("Ward 1","MVGR College Area","18.0552","83.4051"),
    ("Ward 2","Chintalavalasa Junction","18.0675","83.4094"),
    ("Ward 3","RTC Colony","18.0702","83.4153"),
    ("Ward 4","Ramalayam Street","18.0650","83.4005"),
    ("Ward 5","Sai Nagar","18.0751","83.4201")], [1.0,2.5,1.5,1.5])

_h2("3.5 Data Used by the Application")
_body("\u2022  User Data: Registrations, roles, OTP records, Green Points.\n\u2022  Complaint Data: GPS, photos, status history, resolution timestamps.\n\u2022  Schedule Data: Ward-specific collection schedules.\n\u2022  IoT Telemetry: Fill level, battery, temperature, methane.\n\u2022  Waste Declaration Data: Wet, dry, sanitary, hazardous quantities.\n\u2022  Billing Data: PAYT invoices with payment status.")

_h2("3.6 Data Preparation")
_body("A synthetic ML training dataset of 600 rows was generated because sufficient historical telemetry was unavailable. Features include day-of-week, season index, recent complaint volume, and ward identifier. The model is designed to accept real historical data when available.")

_h2("3.7 Database Design")
_body("The system uses PostgreSQL with SQLAlchemy ORM.")
_tbl(["Entity","Key Attributes"],[
    ("User","id, username, email, password_hash, role, phone, green_points"),
    ("Complaint","id, name, phone, ward, description, photo, status, lat, lon, user_id"),
    ("ComplaintStatusLog","id, complaint_id, status, note, created_at"),
    ("SmartBin","id, hardware_id, lat, lon, level, battery, temp, methane, status"),
    ("WorkerProfile","id, user_id, vehicle_id, lat, lon, status, rating"),
    ("WasteDeclaration","id, user_id, wet_kg, dry_kg, sanitary_kg, hazardous_kg, ward"),
    ("PAYTInvoice","id, user_id, period, weight_kg, amount_rs, status"),
    ("PushSubscription","id, user_id, endpoint, p256dh, auth"),
    ("NotificationPreference","id, user_id, complaint_submitted, complaint_assigned, etc.")], [1.8,4.7])

# ════════════════════════════════════════════════════════════════════
# CHAPTER 4 — METHODOLOGY
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("4. METHODOLOGY / SYSTEM DESIGN")

_h2("4.1 Requirement Analysis")
_body("Requirements were categorized into functional (FR) and non-functional (NFR):")
_h3("Functional Requirements")
_tbl(["ID","Module","Description"],[
    ("FR-01","Public Portal","Schedules, anonymous complaints"),
    ("FR-02","Citizen Portal","Track complaints, Green Points"),
    ("FR-03","Admin Portal","Dashboard, worker dispatch, analytics"),
    ("FR-04","Worker Portal","GPS tracking, evidence upload"),
    ("FR-05","IoT Integration","Smart-bin telemetry ingestion"),
    ("FR-06","ML Prediction","Overflow risk prediction"),
    ("FR-07","PAYT Billing","Usage-based invoicing"),
    ("FR-08","Push Notifications","Web push alerts"),
    ("FR-09","Offline Support","PWA with IndexedDB queue"),
    ("FR-10","Bilingual Support","English and Telugu")], [0.8,1.5,4.2])
_h3("Non-Functional Requirements")
_tbl(["ID","Category","Description"],[
    ("NFR-01","Performance","TTFB < 1s, compressed responses"),
    ("NFR-02","Security","OWASP headers, RBAC, bcrypt, OTP"),
    ("NFR-03","Accessibility","WCAG 2.1 AA, ARIA landmarks"),
    ("NFR-04","Scalability","Horizontal via gunicorn workers"),
    ("NFR-05","Offline","PWA with service worker + IndexedDB")], [0.8,1.5,4.2])

_h2("4.2 System Architecture")
_body("The system follows a layered architecture with five distinct layers: Presentation (Browser/PWA with Jinja2 and Bootstrap), Application (Flask route modules), Business Logic (ML engine, job queue, push notifications, PAYT billing), Data (PostgreSQL via SQLAlchemy with Redis caching), and External Services (Open-Meteo, Supabase, Sentry, Twilio).")
_cap("Figure 4.2: Overall System Architecture")
_img(DIAGS['arch'])

_h2("4.3 Development Methodology")
_body("The project followed an iterative approach across six phases: Requirements Gathering (Weeks 1-2), System Design (Weeks 3-4), Core Development (Weeks 5-10), Integration (Weeks 11-14), Enhancement (Weeks 15-18), and Testing/Deployment (Weeks 19-20).")

_h2("4.4 Technology Stack")
_tbl(["Component","Technology","Purpose"],[
    ("Backend","Flask 3.1.3","Web framework"),
    ("Database","PostgreSQL (Supabase)","Relational DB"),
    ("ORM","SQLAlchemy 2.0","Python ORM"),
    ("Frontend","Bootstrap 5 + CSS","Responsive UI"),
    ("ML","scikit-learn 1.9","RandomForest"),
    ("Task Queue","Redis + RQ 2.2","Background jobs"),
    ("WSGI","Gunicorn 26.0","Production server"),
    ("Push","pywebpush 2.0","Web push"),
    ("Security","Flask-Talisman, Limiter","Headers, rate limit"),
    ("Hosting","Render.com","Cloud hosting"),
    ("Docker","Dockerfile","Containerization")], [1.8,2.2,2.5])

_h2("4.5 System Workflow")
_body("The end-to-end workflow: Citizens register → file complaints (GPS + photo) → system validates and stores → admin assigns workers → workers visit location → upload evidence → admin verifies → complaint resolved → citizen notified via push/email → Green Points awarded.")
_cap("Figure 4.4: Complaint Lifecycle Flowchart")
_img(DIAGS['complaint'])

_h2("4.6 Data Flow Diagram")
_body("Primary data flows include: (1) Citizen → Complaint API → Database → Admin Dashboard → Worker Dispatch → Evidence Upload → Resolution; (2) IoT Sensors → Telemetry API → SmartBin Table → Admin Monitor → ML Prediction → Priority Queue; (3) Schedule Request → ML Prediction → Response to Citizen.")

_h2("4.7 Machine Learning Methodology")
_body("The ML module implements a RandomForest regressor trained on a synthetic 600-row dataset. Features: day_of_week, season_idx, recent_complaint_count, ward_id. The model predicts hours_until_90pct_fill for dispatch prioritization. A transparent heuristic fallback ensures the route never errors when the model artifact is unavailable.")
_body("Note: Synthetic data was used during prototype development because sufficient historical telemetry was unavailable. The demonstration validates the prediction pipeline, not real-world accuracy.")
_cap("Figure 4.7: Machine Learning Pipeline")
_img(DIAGS['ml'])

_h2("4.8 Security Architecture")
_body("\u2022  Authentication: Flask-Login with bcrypt + OTP/MFA.\n\u2022  Authorization: RBAC with citizen/worker/admin roles.\n\u2022  Security Headers: All 9 OWASP-recommended headers (CSP, HSTS, X-Content-Type-Options, X-Frame-Options, X-XSS-Protection, Referrer-Policy, Permissions-Policy, COOP, COEP).\n\u2022  Rate Limiting: Flask-Limiter with Redis storage.\n\u2022  Input Validation: Flask-WTF with CSRF protection.\n\u2022  SQL Injection Prevention: SQLAlchemy parameterized queries.")

_h2("4.9 PWA and Offline Methodology")
_body("\u2022  Service Worker: Versioned precache manifest for offline support.\n\u2022  IndexedDB Queue: Complaints stored offline, synced on reconnection.\n\u2022  Background Sync: Automatic submission when connectivity returns.\n\u2022  Web App Manifest: Installable with shortcuts, screenshots, standalone mode.\n\u2022  Splash Screen: Branded loading animation on PWA install.\n\u2022  Install Banner: Mobile prompt after second visit with iOS fallback.")
_cap("Figure 4.9: PWA Offline Workflow")
_img(DIAGS['pwa'])

# ════════════════════════════════════════════════════════════════════
# CHAPTER 5 — IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("5. IMPLEMENTATION / MODULES")
modules = [
    ("5.1 Public Portal","Provides waste-management information to all visitors.",
     "Schedule display, anonymous complaints, tracking, ward transparency, FAQ, search, RSS, AI-friendly llms.txt",
     "Flask + Jinja2 + Bootstrap 5","Fully implemented"),
    ("5.2 Citizen Portal","Authenticated portal for citizens.",
     "Registration with OTP, complaint filing, dashboard, Green Points, waste declarations, push notifications",
     "Flask-Login + SQLAlchemy + Socket.IO","Fully implemented"),
    ("5.3 Admin Portal","Comprehensive admin dashboard.",
     "Complaint overview, worker dispatch, IoT monitoring, ML display, PAYT management, push analytics, audit logs",
     "Flask-Login + RBAC + Socket.IO","Fully implemented"),
    ("5.4 Worker Portal","Mobile-friendly worker portal.",
     "Dispatch acceptance, GPS tracking, photo evidence, task queue, performance rating",
     "Flask-Login + Geolocation API","Fully implemented"),
    ("5.5 IoT Smart Bin Module","Telemetry ingestion and monitoring.",
     "Authenticated API, fill level, battery, temperature, methane, status classification",
     "Flask API + SQLAlchemy","Implemented (simulated)"),
    ("5.6 Machine Learning Module","Overflow risk prediction.",
     "RandomForest regressor, feature engineering, transparent fallback, pickle persistence",
     "scikit-learn + pandas + numpy","Implemented (synthetic data)"),
    ("5.7 Green Points Module","Gamification for segregation.",
     "Points earned, streak tracking, leaderboard, redemption",
     "Flask-Login + SQLAlchemy","Fully implemented"),
    ("5.8 PAYT Module","Usage-based billing.",
     "Invoice generation, compliance scoring, UPI/Razorpay, PDF via ReportLab",
     "Flask + ReportLab + Razorpay","Implemented (test mode)"),
    ("5.9 Background Jobs","Async task processing.",
     "Status notifications, SLA escalation, dunning, email, push dispatch",
     "Redis + RQ","Fully implemented"),
    ("5.10 PWA and Offline","Offline capabilities.",
     "Service worker, IndexedDB queue, background sync, manifest, splash screen, install banner",
     "SW API + IndexedDB","Fully implemented"),
]
for heading, purpose, functions, tech, status in modules:
    _h2(heading)
    _h3("Purpose"); _body(purpose)
    _h3("Key Functions"); _body(functions)
    _h3("Technology"); _body(tech)
    _h3("Status"); _body(status)

# ════════════════════════════════════════════════════════════════════
# CHAPTER 6 — RESULTS
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("6. RESULTS / OUTPUTS")
_h2("6.1 Implemented Features Summary")
_tbl(["Module","Features","Status","Verification"],[
    ("Public Portal","Homepage, schedule, transparency, search","Fully Implemented","200 OK"),
    ("Citizen Portal","Registration, complaints, dashboard, Green Points","Fully Implemented","200 OK"),
    ("Admin Portal","Dashboard, dispatch, IoT, ML display","Fully Implemented","200 OK"),
    ("Worker Portal","Dispatch, GPS, evidence upload","Fully Implemented","200 OK"),
    ("IoT Integration","Telemetry API, sensor monitoring","Simulated","API functional"),
    ("ML Prediction","RandomForest overflow prediction","Synthetic Data","Pipeline functional"),
    ("PAYT Billing","Invoice generation, payment","Test Mode","Invoice OK"),
    ("Push Notifications","Web push, preferences, analytics","Fully Implemented","API functional"),
    ("PWA / Offline","Service worker, IndexedDB, manifest","Fully Implemented","All 200"),
    ("Security","OWASP headers, RBAC, bcrypt, OTP","Fully Implemented","9/9 headers"),
    ("Accessibility","ARIA, skip-to-content, contrast","Implemented","Checks pass")], [1.3,2.5,1.5,1.2])

_h2("6.2 Homepage Output")
_body("The homepage displays collection schedule lookup, complaint filing shortcut, ward transparency map, weather widget, and community impact statistics.")
_cap("Figure 6.1: Homepage — Collection schedule, impact stats, weather widget")

_h2("6.3 Collection Schedule Output")
_body("The schedule page shows ward-specific collection timetables with ML overflow prediction.")
_cap("Figure 6.2: Collection Schedule — Ward selection with ML prediction")

_h2("6.4 Complaint Reporting Output")
_body("The complaint form captures name, phone, ward, address, description, photo, and automatic GPS coordinates.")
_cap("Figure 6.3: Complaint Reporting — GPS-enabled form with photo upload")

_h2("6.5 Citizen Dashboard Output")
_body("The citizen dashboard shows complaint history, Green Points, waste declarations, and notification preferences.")
_cap("Figure 6.4: Citizen Dashboard — History, Green Points, quick actions")

_h2("6.6 Admin Dashboard Output")
_body("The admin control room displays real-time complaints, IoT bin status, worker queue, and push analytics.")
_cap("Figure 6.5: Admin Dashboard — Real-time monitoring with IoT and analytics")

_h2("6.7 Performance Evaluation")
_tbl(["Metric","Value","Notes"],[
    ("Homepage TTFB","~0.5s","Gunicorn sync worker"),
    ("Static Cache","31,536,000s","Immutable, version-busted"),
    ("HTML Cache (repeat)","60s","Browser + ETag 304"),
    ("Brotli Compression","Level 4","All text responses"),
    ("Preload Hints","4 Link headers","Critical resources"),
    ("FCP","~0.5s","With preload hints")], [2.0,1.8,2.7])

_h2("6.8 Security Evaluation")
_tbl(["Header","Status","Purpose"],[
    ("Content-Security-Policy","Present","Restricts resource loading"),
    ("Strict-Transport-Security","Present","Enforces HTTPS"),
    ("X-Content-Type-Options","Present","nosniff"),
    ("X-Frame-Options","Present","DENY clickjacking"),
    ("X-XSS-Protection","Present","Legacy XSS protection"),
    ("Referrer-Policy","Present","strict-origin-when-cross-origin"),
    ("Permissions-Policy","Present","Restricts browser features"),
    ("Cross-Origin-Opener-Policy","Present","same-origin isolation"),
    ("Cross-Origin-Embedder-Policy","Present","require-corp")], [2.5,1.0,3.0])

_h2("6.9 Accessibility Evaluation")
_body("\u2022  Skip-to-content link for keyboard navigation.\n\u2022  ARIA landmarks (navigation, main, contentinfo) on all pages.\n\u2022  Semantic HTML5 elements.\n\u2022  Sufficient color contrast (minimum 4.5:1).\n\u2022  Form labels associated with inputs.\n\u2022  Keyboard-navigable elements.\n\u2022  Alt text for images and icons.")

# ════════════════════════════════════════════════════════════════════
# CHAPTER 7
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("7. IMPACT ASSESSMENT")
_body("Note: As a prototype, impact values are estimated projections based on system capabilities, not measured community pilot results.")
_h2("7.1 Social Impact")
_body("\u2022  Citizen Empowerment: Transparent access to schedules, complaints, and performance metrics.\n\u2022  Accessibility: Bilingual support and PWA offline for low-connectivity areas.\n\u2022  Accountability: Digital tracking with status timelines and SLA monitoring.")
_h2("7.2 Operational Impact")
_tbl(["Metric","With SmartGarbage","Without"],[
    ("Schedule Access","100% digital","Verbal/informal"),
    ("Complaint Tracking","Real-time","Untracked"),
    ("Worker Dispatch","GPS-tracked","Manual"),
    ("Data Decisions","ML + analytics","Experience-based")], [2.0,2.2,2.3])
_h2("7.3 Environmental Impact")
_body("\u2022  Green Points incentivize source segregation, potentially increasing recycling.\n\u2022  IoT monitoring prevents overflow and environmental contamination.\n\u2022  Estimated CO2 savings: 0.5 kg CO2 per kg recycled material (EPA estimate).")
_h2("7.4 Economic Feasibility")
_body("\u2022  Zero-cost deployment on Render.com free tier + Supabase free PostgreSQL.\n\u2022  PAYT billing generates revenue while incentivizing segregation.\n\u2022  Optimized routes reduce fuel and labor costs.\n\u2022  Open source eliminates vendor lock-in.")
_h2("7.5 Scalability")
_body("Architecture supports horizontal scaling via gunicorn workers. Multi-panchayat deployment achievable by extending WARD_COORDINATES mapping with ward-based database isolation.")

# ════════════════════════════════════════════════════════════════════
# CHAPTER 8
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("8. CHALLENGES FACED")
challenges = [
    ("Data Unavailability","Insufficient historical telemetry for ML","Generated 600-row synthetic dataset; documented limitation for future work"),
    ("Offline Architecture","Complaint filing without internet","Service Worker + IndexedDB + Background Sync API"),
    ("Real-time Updates","Live dashboard without refresh","Flask-SocketIO with Redis broker"),
    ("IoT Simulation","No physical hardware","Authenticated API with test data generators"),
    ("Security Hardening","Comprehensive security without usability loss","All 9 OWASP headers via Flask-Talisman"),
    ("Performance","TTFB on free-tier hosting","Multi-layer caching, Brotli, preload hints"),
    ("Bilingual Support","Consistent translations","Flask-Babel with gettext markers"),
    ("Push Notifications","Reliable delivery","pywebpush with VAPID + preference filtering"),
    ("Deployment","Dev/prod consistency","Docker containerization + env vars + health check"),
    ("Payment Integration","Testing without live merchant","Razorpay test keys + UPI primary method")]
for challenge, problem, solution in challenges:
    _h3(challenge)
    _body(f"Problem: {problem}\nSolution: {solution}")

# ════════════════════════════════════════════════════════════════════
# CHAPTER 9
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("9. CONCLUSION")
_h2("9.1 Summary")
_body("SmartGarbage Chintalavalasa is a community-based smart waste management and digital governance system integrating Flask, PostgreSQL, IoT telemetry, ML prediction, PWA offline capabilities, PAYT billing, and comprehensive security into a unified platform.")
_h2("9.2 Achievement of Objectives")
_body("The project successfully achieved all eight stated objectives: web platform, admin dashboard, IoT integration, ML prediction, PAYT billing, PWA capabilities, security implementation, and bilingual support.")
_h2("9.3 Limitations")
_body("\u2022  ML trained on synthetic data, not validated with real telemetry.\n\u2022  IoT telemetry simulated without physical sensors.\n\u2022  PAYT in test mode without live merchant account.\n\u2022  Impact values are estimated, not measured from a pilot.\n\u2022  Not tested with real users at community scale.\n\u2022  SMS/WhatsApp require third-party API keys.")

# ════════════════════════════════════════════════════════════════════
# CHAPTER 10
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("10. FUTURE WORK")
for section, items in [
    ("10.1 Real-World Data Integration",["Replace synthetic ML data with real telemetry","Conduct community pilot with actual residents","Collect baseline data for before/after studies"]),
    ("10.2 IoT Hardware Deployment",["Deploy ultrasonic fill-level sensors","GPS trackers on collection vehicles","LoRaWAN or NB-IoT connectivity"]),
    ("10.3 Mobile Application",["Native Android application","WhatsApp Business API chatbot","SMS schedule notifications"]),
    ("10.4 Advanced Analytics",["Time-series forecasting","Route optimization","Computer vision waste classification"]),
    ("10.5 Multi-Panchayat Deployment",["Tenant isolation for multiple panchayats","District-level aggregation","SBM Grameen reporting compliance"]),
    ("10.6 Payment and Billing",["Live Razorpay merchant account","Recurring billing via UPI mandates","Payment analytics and revenue reporting"])]:
    _h2(section)
    for item in items: _p(f"\u2022  {item}", 12, sa=4)

# ════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("REFERENCES"); doc.add_paragraph()
refs = [
    '[1] GOV.UK, "Design System," Government Digital Service, 2024. https://design-system.service.gov.uk/',
    '[2] Ministry of Housing and Urban Affairs, "Swachh Bharat Mission — Urban," Government of India, 2024.',
    '[3] Ministry of Jal Shakti, "Swachh Bharat Mission — Grameen Phase II," Government of India, 2024.',
    '[4] T. Anagnostopoulos et al., "Waste Management as an IoT-Enabled Service in Smart Cities," Springer, 2015.',
    '[5] M. A. ad Din et al., "Smart Bin: IoT-Based Waste Monitoring System," Proc. IT, 2020.',
    '[6] S. Kumar and R. Sharma, "IoT-Based Smart Waste Management: A Review," J. Cleaner Production, vol. 295, 2021.',
    '[7] M. Afshin et al., "ML Approaches for MSW Generation Prediction," Waste Management, vol. 125, 2021.',
    '[8] Y. Chen et al., "Predicting MSW Using ML Methods," Env. Sci. Pollution Research, vol. 27, 2020.',
    '[9] D. Thung and M. Yang, "Waste Classification using CNN," AISI, Springer, 2016.',
    '[10] W3C, "WCAG 2.1," World Wide Web Consortium, 2018.',
    '[11] OWASP Foundation, "OWASP Top Ten 2021," 2021.',
    '[12] U.S. EPA, "Pay-As-You-Throw," 2023.',
    '[13] Flask Documentation, "Flask — Pallets Projects," 2024.',
    '[14] SQLAlchemy Documentation, "SQLAlchemy," 2024.',
    '[15] scikit-learn Documentation, "scikit-learn," 2024.',
    '[16] MDN Web Docs, "Progressive Web Apps," Mozilla, 2024.',
    '[17] Web Push Protocol, "Push API — MDN," Mozilla, 2024.',
    '[18] Supabase Documentation, "Supabase," 2024.',
    '[19] Render Documentation, "Render," 2024.',
    '[20] Razorpay Documentation, "Razorpay Docs," 2024.']
for r in refs: _p(r, 12, sa=6)

# ════════════════════════════════════════════════════════════════════
# APPENDIX A
# ════════════════════════════════════════════════════════════════════
_pb(); _h1("APPENDIX A: PACKAGES, TOOLS USED & WORKING PROCESS")
_h2("A.1 Packages and Tools")
_tbl(["Package","Category","Purpose"],[
    ("Flask 3.1.3","Python","Web framework"),("Flask-SQLAlchemy 3.1.1","Python","ORM"),
    ("Flask-Login 0.6.3","Python","Sessions"),("Flask-Talisman 1.1.0","Python","Security headers"),
    ("Flask-Limiter 4.1.1","Python","Rate limiting"),("Flask-Compress 1.17","Python","Compression"),
    ("Flask-WTF 1.2.1","Python","Forms + CSRF"),("Flask-SocketIO 5.3.6","Python","Real-time WebSocket"),
    ("SQLAlchemy 2.0.50","Python","ORM"),("psycopg2-binary 2.9","Python","PostgreSQL"),
    ("Redis 6.2.0","Python","Cache client"),("RQ 2.2.0","Python","Job queue"),
    ("scikit-learn 1.9.0","Python","ML"),("pandas 3.0.3","Python","Data processing"),
    ("numpy 2.4.6","Python","Numerical"),("pywebpush 2.0.0","Python","Push notifications"),
    ("reportlab 5.0.0","Python","PDF generation"),("sentry-sdk 2.5.0","Python","Error tracking"),
    ("gunicorn 26.0.0","Python","WSGI server"),("Jinja2 3.1.6","Python","Templates"),
    ("Bootstrap 5","CSS/JS","UI framework"),("Socket.IO","JS","Real-time client"),
    ("Docker","DevOps","Containerization"),("Git","DevOps","Version control"),
    ("PostgreSQL (Supabase)","Database","Relational DB"),("Redis","Database","In-memory store"),
    ("Render.com","Cloud","Hosting"),("Sentry","Monitoring","Error tracking")], [2.5,1.2,2.8])
_h2("A.2 Working Process")
_body("1. Environment Setup: Python 3.12 virtual env, Flask dependencies, PostgreSQL on Supabase, Redis on Upstash.\n2. Project Structure: app/routes/, app/templates/, app/static/, app/models.py.\n3. Version Control: Git with GitHub repository.\n4. Iterative Development: Public → Citizen → Admin → Worker → IoT → ML → PWA.\n5. Testing: Flask test client, python -m compileall, curl security checks.\n6. Deployment: Docker multi-stage build, Render.com auto-deploy from GitHub.\n7. Monitoring: Sentry integration, /health endpoint, structured logging.")

# APPENDIX B
_pb(); _h1("APPENDIX B: IMPORTANT SOURCE CODE")
_h2("B.1 Application Factory (app/__init__.py)")
_body("Flask application factory initializing extensions, security headers, caching, and background jobs.")
_h2("B.2 Database Models (app/models.py)")
_body("SQLAlchemy models: User, Complaint, SmartBin, WorkerProfile, WasteDeclaration, PAYTInvoice, PushSubscription, NotificationPreference, ConsentRecord.")
_h2("B.3 ML Module (app/ml_model.py)")
_body("RandomForest regressor with synthetic data training and transparent heuristic fallback.")
_h2("B.4 Service Worker (app/static/sw.js)")
_body("Versioned precache manifest, stale-while-revalidate for HTML, cache-first for assets, push event handler.")
_h2("B.5 Push Module (app/push.py)")
_body("VAPID key management, preference-based filtering, delivery logging, complaint lifecycle integration.")

# PAPER PUBLICATIONS
_pb(); _h1("PAPER PUBLICATIONS"); doc.add_paragraph()
_body("No research papers have been published based on this project at the time of report submission.")

# ════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SmartGarbage_Community_Project_Report.docx")
doc.save(out)
sz = os.path.getsize(out)
print(f"\nReport generated: {out}")
print(f"Size: {sz:,} bytes ({sz/1024:.1f} KB)")

# Verify images
vdoc = Document(out)
imgs = sum(1 for r in vdoc.part.rels.values() if 'image' in r.reltype)
print(f"Images embedded: {imgs}")
print(f"Total paragraphs: {len(vdoc.paragraphs)}")
print(f"Total tables: {len(vdoc.tables)}")
