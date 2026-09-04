#!/usr/bin/env python3
"""
Create 4 professional academic diagrams and embed them into the DOCX report.
Diagrams: System Architecture, Complaint Lifecycle, ML Pipeline, PWA Workflow.
"""

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
import os

OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_diagrams")
os.makedirs(OUT_DIR, exist_ok=True)

# ── Color palette ───────────────────────────────────────────────────
C = {
    'dark_green': '#1B5E20',
    'green': '#388E3C',
    'light_green': '#C8E6C9',
    'pale_green': '#E8F5E9',
    'blue': '#1565C0',
    'light_blue': '#BBDEFB',
    'pale_blue': '#E3F2FD',
    'orange': '#E65100',
    'light_orange': '#FFE0B2',
    'purple': '#6A1B9A',
    'light_purple': '#E1BEE7',
    'red': '#B71C1C',
    'light_red': '#FFCDD2',
    'gray': '#424242',
    'light_gray': '#E0E0E0',
    'white': '#FFFFFF',
    'black': '#000000',
}

def save(fig, name):
    path = os.path.join(OUT_DIR, name)
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(fig)
    print(f"  Saved: {path}")
    return path

def box(ax, x, y, w, h, text, color, text_color='white', fontsize=9, bold=False, radius=0.02):
    """Draw a rounded rectangle with centered text."""
    fancy = FancyBboxPatch((x - w/2, y - h/2), w, h,
                           boxstyle=f"round,pad=0.01,rounding_size={radius}",
                           facecolor=color, edgecolor=C['gray'], linewidth=1.2)
    ax.add_patch(fancy)
    weight = 'bold' if bold else 'normal'
    ax.text(x, y, text, ha='center', va='center', fontsize=fontsize,
            fontweight=weight, color=text_color, wrap=True,
            multialignment='center')

def arrow(ax, x1, y1, x2, y2, color=C['gray'], style='->', lw=1.5):
    """Draw an arrow between two points."""
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle=style, color=color, lw=lw))


# ════════════════════════════════════════════════════════════════════
# DIAGRAM 1: System Architecture
# ════════════════════════════════════════════════════════════════════
def create_system_architecture():
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7.5)
    ax.axis('off')
    ax.set_title('Figure 4.2: Overall System Architecture', fontsize=13, fontweight='bold', pad=15)

    # Layer 1: Users
    box(ax, 5, 7.0, 8.0, 0.6, 'USERS', C['dark_green'], fontsize=12, bold=True)
    users = ['Citizen', 'Admin', 'Worker', 'Public']
    for i, u in enumerate(users):
        box(ax, 1.5 + i*2.3, 7.0, 1.6, 0.4, u, C['green'], fontsize=8)

    arrow(ax, 5, 6.6, 5, 6.15)

    # Layer 2: Browser/PWA
    box(ax, 5, 6.0, 8.0, 0.55, 'BROWSER / MOBILE / PWA', C['blue'], fontsize=10, bold=True)

    arrow(ax, 5, 5.65, 5, 5.2)

    # Layer 3: CDN
    box(ax, 5, 5.05, 5.0, 0.4, 'CDN / Service Worker Cache', C['light_blue'], text_color=C['blue'], fontsize=9, bold=True)

    arrow(ax, 5, 4.78, 5, 4.35)

    # Layer 4: Flask Application
    box(ax, 5, 4.15, 8.5, 0.55, 'FLASK APPLICATION (Gunicorn + Greenlet)', C['orange'], fontsize=10, bold=True)

    arrow(ax, 5, 3.78, 5, 3.35)

    # Layer 5: Portal Modules
    box(ax, 5, 3.1, 8.5, 0.5, '', C['light_orange'], text_color=C['orange'])
    modules = ['Public\nPortal', 'Citizen\nPortal', 'Admin\nPortal', 'Worker\nPortal', 'IoT\nAPI']
    for i, m in enumerate(modules):
        box(ax, 0.85 + i*1.8, 3.1, 1.5, 0.45, m, C['orange'], fontsize=7, bold=True)

    arrow(ax, 5, 2.75, 5, 2.3)

    # Layer 6: Business Logic
    box(ax, 5, 2.1, 8.5, 0.5, '', C['light_purple'], text_color=C['purple'])
    logic = ['ML Engine', 'Job Queue\n(RQ)', 'Push\nNotif.', 'PAYT\nBilling']
    for i, l in enumerate(logic):
        box(ax, 1.3 + i*2.3, 2.1, 1.8, 0.45, l, C['purple'], fontsize=7, bold=True)

    arrow(ax, 5, 1.75, 5, 1.35)

    # Layer 7: Data
    box(ax, 5, 1.15, 8.5, 0.5, '', C['light_green'], text_color=C['green'])
    data_items = ['PostgreSQL\n(Supabase)', 'Redis', 'Sentry']
    for i, d in enumerate(data_items):
        box(ax, 1.8 + i*2.8, 1.15, 2.0, 0.45, d, C['green'], fontsize=7, bold=True)

    arrow(ax, 5, 0.82, 5, 0.4)

    # Layer 8: External
    box(ax, 5, 0.25, 8.5, 0.45, 'EXTERNAL SERVICES:  Open-Meteo API  |  Twilio  |  Razorpay  |  Render.com  |  GitHub',
        C['light_gray'], text_color=C['gray'], fontsize=8, bold=True)

    return save(fig, 'system_architecture.png')


# ════════════════════════════════════════════════════════════════════
# DIAGRAM 2: Complaint Lifecycle
# ════════════════════════════════════════════════════════════════════
def create_complaint_lifecycle():
    fig, ax = plt.subplots(figsize=(10, 8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_title('Figure 4.4: Complaint Lifecycle Flowchart', fontsize=13, fontweight='bold', pad=15)

    # Start
    box(ax, 5, 8.5, 4.0, 0.55, 'Citizen Reports Issue\n(GPS + Photo + Description)', C['green'], fontsize=9, bold=True)

    arrow(ax, 5, 8.15, 5, 7.7)

    # Validation
    box(ax, 5, 7.5, 3.5, 0.5, 'Server-side Validation', C['blue'], fontsize=9, bold=True)

    arrow(ax, 5, 7.18, 5, 6.75)

    # Duplicate check
    box(ax, 5, 6.5, 3.5, 0.5, 'Duplicate Complaint Detection', C['blue'], fontsize=9, bold=True)

    # Branch
    arrow(ax, 3.5, 6.25, 2.0, 5.75)
    arrow(ax, 6.5, 6.25, 8.0, 5.75)

    # Duplicate
    box(ax, 2.0, 5.5, 2.5, 0.5, 'Duplicate\n→ Notify Citizen', C['light_red'], text_color=C['red'], fontsize=8)

    # Valid
    box(ax, 8.0, 5.5, 2.5, 0.5, 'Valid\n→ Create Complaint', C['light_green'], text_color=C['green'], fontsize=8, bold=True)

    arrow(ax, 8.0, 5.18, 8.0, 4.75)

    # Admin assignment
    box(ax, 8.0, 4.5, 2.8, 0.5, 'Admin Reviews\n& Assigns Worker', C['orange'], fontsize=8, bold=True)

    arrow(ax, 8.0, 4.18, 5.0, 3.6)

    # Worker dispatch
    box(ax, 5, 3.35, 3.0, 0.5, 'Worker Dispatched\n(GPS Tracked)', C['purple'], fontsize=9, bold=True)

    arrow(ax, 5, 3.0, 5, 2.55)

    # Evidence
    box(ax, 5, 2.3, 3.5, 0.5, 'Worker Uploads\nPhoto + GPS Evidence', C['blue'], fontsize=9, bold=True)

    arrow(ax, 5, 2.0, 5, 1.55)

    # Verify
    box(ax, 5, 1.3, 3.5, 0.5, 'Admin Verifies\nResolution', C['orange'], fontsize=9, bold=True)

    arrow(ax, 5, 1.0, 5, 0.55)

    # Resolved
    box(ax, 5, 0.3, 3.5, 0.5, '✓ Complaint Resolved\n+ Citizen Notified + Green Points', C['green'], fontsize=8, bold=True)

    # Side note: notification
    box(ax, 1.5, 4.5, 2.2, 0.55, 'Push / Email\nNotification', C['light_blue'], text_color=C['blue'], fontsize=7)
    arrow(ax, 3.2, 4.5, 3.8, 3.35, color=C['blue'], style='->')

    return save(fig, 'complaint_lifecycle.png')


# ════════════════════════════════════════════════════════════════════
# DIAGRAM 3: ML Pipeline
# ════════════════════════════════════════════════════════════════════
def create_ml_pipeline():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.5)
    ax.axis('off')
    ax.set_title('Figure 4.7: Machine Learning Pipeline', fontsize=13, fontweight='bold', pad=15)

    # Pipeline stages (left to right)
    stages = [
        (1.0, 4.5, 'Available\nPrototype Data', C['light_gray'], C['gray']),
        (2.8, 4.5, 'Data\nPreparation', C['blue'], 'white'),
        (4.6, 4.5, 'Feature\nEngineering', C['blue'], 'white'),
        (6.4, 4.5, 'Synthetic\nTraining Set\n(600 rows)', C['orange'], 'white'),
        (8.2, 4.5, 'RandomForest\nModel\nTraining', C['purple'], 'white'),
    ]

    for x, y, text, color, tc in stages:
        box(ax, x, y, 1.6, 1.0, text, color, text_color=tc, fontsize=7, bold=True)
    for i in range(len(stages)-1):
        arrow(ax, stages[i][0]+0.8, 4.5, stages[i+1][0]-0.8, 4.5)

    # Bottom row: prediction flow
    stages2 = [
        (8.2, 2.5, 'Trained Model\n(Pickle)', C['purple'], 'white'),
        (6.0, 2.5, 'Prediction:\nHours Until\n90% Fill', C['green'], 'white'),
        (3.8, 2.5, 'Rank Bins\nby Urgency', C['orange'], 'white'),
        (1.6, 2.5, 'Priority\nDispatch', C['dark_green'], 'white'),
    ]
    for x, y, text, color, tc in stages2:
        box(ax, x, y, 1.6, 1.0, text, color, text_color=tc, fontsize=7, bold=True)
    for i in range(len(stages2)-1):
        arrow(ax, stages2[i][0]-0.8, 2.5, stages2[i+1][0]+0.8, 2.5)

    # Connect top right to bottom right
    arrow(ax, 8.2, 3.95, 8.2, 3.1)

    # Fallback note
    box(ax, 3.0, 1.0, 5.5, 0.55, 'Fallback: Transparent heuristic when model unavailable',
        C['light_blue'], text_color=C['blue'], fontsize=8, bold=False)
    arrow(ax, 1.6, 2.15, 3.0, 1.3, color=C['blue'], style='->')

    # Note
    ax.text(5, 0.3, 'Note: Synthetic data used during prototype development.\n'
            'Real historical telemetry integration proposed for future work.',
            ha='center', va='center', fontsize=8, fontstyle='italic', color=C['gray'])

    return save(fig, 'ml_pipeline.png')


# ════════════════════════════════════════════════════════════════════
# DIAGRAM 4: PWA Offline Workflow
# ════════════════════════════════════════════════════════════════════
def create_pwa_workflow():
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.5)
    ax.axis('off')
    ax.set_title('Figure 4.9: PWA Offline Workflow', fontsize=13, fontweight='bold', pad=15)

    # Start
    box(ax, 5, 6.0, 3.5, 0.55, 'User Submits Complaint', C['green'], fontsize=10, bold=True)

    arrow(ax, 5, 5.65, 5, 5.2)

    # Decision
    # Diamond shape using text
    diamond_x, diamond_y = 5, 4.9
    diamond = plt.Polygon([(5, 5.2), (6.3, 4.9), (5, 4.6), (3.7, 4.9)],
                          facecolor=C['light_orange'], edgecolor=C['orange'], linewidth=1.5)
    ax.add_patch(diamond)
    ax.text(5, 4.9, 'Internet\nAvailable?', ha='center', va='center', fontsize=8, fontweight='bold')

    # Yes branch (left)
    arrow(ax, 3.7, 4.9, 2.0, 4.9, color=C['green'], style='->')
    ax.text(2.8, 5.05, 'Yes', ha='center', va='center', fontsize=8, color=C['green'], fontweight='bold')
    box(ax, 2.0, 4.5, 2.2, 0.55, 'Submit via\nAPI → Database', C['green'], fontsize=8, bold=True)

    arrow(ax, 2.0, 4.15, 2.0, 3.65)
    box(ax, 2.0, 3.4, 2.2, 0.5, 'Success ✓\nCitizen Notified', C['light_green'], text_color=C['green'], fontsize=8)

    # No branch (right)
    arrow(ax, 6.3, 4.9, 8.0, 4.9, color=C['red'], style='->')
    ax.text(7.2, 5.05, 'No', ha='center', va='center', fontsize=8, color=C['red'], fontweight='bold')
    box(ax, 8.0, 4.5, 2.2, 0.55, 'Store in\nIndexedDB Queue', C['orange'], fontsize=8, bold=True)

    arrow(ax, 8.0, 4.15, 8.0, 3.65)
    box(ax, 8.0, 3.4, 2.2, 0.5, 'Background\nSync Registered', C['light_orange'], text_color=C['orange'], fontsize=8)

    arrow(ax, 8.0, 3.08, 5.5, 2.5)
    ax.text(6.8, 2.75, 'Connection\nRestored', ha='center', va='center', fontsize=7, color=C['gray'])

    # Sync step
    box(ax, 5, 2.2, 3.0, 0.55, 'Service Worker\nSync Event Fires', C['blue'], fontsize=8, bold=True)

    arrow(ax, 5, 1.85, 5, 1.35)
    box(ax, 5, 1.1, 3.5, 0.55, 'API Submission\n→ Queue Emptied', C['green'], fontsize=8, bold=True)

    arrow(ax, 5, 0.75, 5, 0.3)
    box(ax, 5, 0.1, 3.0, 0.4, 'Citizen Notified ✓', C['dark_green'], fontsize=8, bold=True)

    return save(fig, 'pwa_workflow.png')


# ════════════════════════════════════════════════════════════════════
# MAIN: Create all diagrams and embed into DOCX
# ════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print("Creating diagrams...")
    p1 = create_system_architecture()
    p2 = create_complaint_lifecycle()
    p3 = create_ml_pipeline()
    p4 = create_pwa_workflow()

    print("\nEmbedding diagrams into DOCX...")
    from docx import Document
    from docx.shared import Inches

    docx_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SmartGarbage_Community_Project_Report.docx")
    doc = Document(docx_path)

    # Helper: find a paragraph by text substring
    def find_para_index(text_substring, start=0):
        for i, p in enumerate(doc.paragraphs):
            if i >= start and text_substring in p.text:
                return i
        return None

    # Helper: insert image after a paragraph index
    def insert_image_after(para_idx, image_path, width_inches=5.5):
        p = doc.paragraphs[para_idx]
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
        return para_idx + 1

    diagrams = [
        ('system_architecture.png', 'Figure 4.2: Overall System Architecture', 'System Architecture'),
        ('complaint_lifecycle.png', 'Figure 4.4: Complaint Lifecycle Flowchart', 'Complaint Lifecycle'),
        ('ml_pipeline.png', 'Figure 4.7: Machine Learning Pipeline', 'ML Pipeline'),
        ('pwa_workflow.png', 'Figure 4.9: PWA Offline Workflow', 'PWA Workflow'),
    ]

    # Strategy: find the figure captions and insert images before them
    for fname, caption, label in diagrams:
        img_path = os.path.join(OUT_DIR, fname)
        # Find the caption paragraph
        idx = find_para_index(caption)
        if idx is not None:
            # Insert the image BEFORE the caption paragraph
            p = doc.paragraphs[idx]
            run = p.add_run()
            run.add_picture(img_path, width=Inches(5.5))
            # Move image to before the caption by restructuring XML
            # Actually, we need to insert BEFORE. Let's use a different approach.
            print(f"  Found caption for {label} at para {idx}")
        else:
            print(f"  WARNING: Caption '{caption}' not found for {label}")

    # Better approach: insert images before caption paragraphs by modifying XML
    doc2 = Document(docx_path)
    inserted = 0

    for fname, caption, label in diagrams:
        img_path = os.path.join(OUT_DIR, fname)
        idx = find_para_index(caption)
        if idx is None:
            continue

        # Create a new paragraph element with the image, insert before caption
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Get the parent element (body)
        body = doc2.paragraphs[idx]._element.getparent()

        # Create new paragraph
        new_p = OxmlElement('w:p')
        new_r = OxmlElement('w:r')

        # Add the image
        from docx.shared import Emu
        img_width_emu = int(Inches(5.5))

        # Create relationship
        rel = doc2.part.get_or_add_image_relationship(img_path)

        # Build the drawing element
        drawing = OxmlElement('w:drawing')
        inline = OxmlElement('wp:inline')

        extent = OxmlElement('wp:extent')
        extent.set('cx', str(img_width_emu))
        extent.set('cy', str(int(img_width_emu * 0.7)))
        inline.append(extent)

        docPr = OxmlElement('wp:docPr')
        docPr.set('id', str(inserted + 1))
        docPr.set('name', label)
        inline.append(docPr)

        blipFill = OxmlElement('a:blipFill')
        blip = OxmlElement('a:blip')
        blip.set(qn('r:embed'), rel.rId)
        blipFill.append(blip)
        inline.append(blipFill)

        # Graphic frame
        graphic = OxmlElement('a:graphic')
        graphicFrame = OxmlElement('a:graphicFrame')
        xfrm = OxmlElement('a:xfrm')
        off = OxmlElement('a:off')
        off.set('x', '0')
        off.set('y', '0')
        ext = OxmlElement('a:ext')
        ext.set('cx', str(img_width_emu))
        ext.set('cy', str(int(img_width_emu * 0.7)))
        xfrm.append(off)
        xfrm.append(ext)
        graphicFrame.append(xfrm)
        graphic.append(graphicFrame)
        inline.append(graphic)

        drawing.append(inline)
        new_r.append(drawing)
        new_p.append(new_r)

        # Insert before the caption paragraph
        body.insert(list(body).index(doc2.paragraphs[idx]._element), new_p)
        inserted += 1
        print(f"  Embedded: {label}")

    doc2.save(docx_path)
    print(f"\nDone! {inserted} diagrams embedded into {docx_path}")

    # Verify
    verify_doc = Document(docx_path)
    img_count = 0
    for rel in verify_doc.part.rels.values():
        if 'image' in rel.reltype:
            img_count += 1
    print(f"Verification: {img_count} images found in document")
