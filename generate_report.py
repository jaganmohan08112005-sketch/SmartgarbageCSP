#!/usr/bin/env python3
"""
Generate the B.Tech Community Project Report for SmartGarbage Chintalavalasa.
Produces a professionally formatted .DOCX file following MVGR College guidelines.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page setup ──────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)

FONT_NAME = "Times New Roman"

# ── Style helpers ───────────────────────────────────────────────────
def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT_NAME}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), FONT_NAME)

def add_paragraph(text, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, space_before=0, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_heading_chapter(text):
    """Main chapter heading: 16pt, Bold, ALL CAPS"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.0
    run = p.add_run(text.upper())
    set_font(run, size=16, bold=True)
    return p

def add_heading_section(text):
    """Section heading: 14pt Bold"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    return p

def add_heading_subsection(text):
    """Subsection heading: 12pt Bold"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    run = p.add_run(text)
    set_font(run, size=12, bold=True)
    return p

def add_body(text):
    """Body text: 12pt, justified"""
    return add_paragraph(text, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

def add_body_bold(text):
    return add_paragraph(text, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

def add_figure_caption(text):
    """Figure/table caption: 12pt bold centered"""
    return add_paragraph(text, size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

def add_page_break():
    doc.add_page_break()

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_table(headers, rows, col_widths=None):
    """Add a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=12, bold=True)
        set_cell_shading(cell, "D9E2F3")
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            set_font(run, size=12)
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()  # spacing after table
    return table


# ════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════════
for _ in range(3):
    doc.add_paragraph()

add_paragraph("COMMUNITY PROJECT REPORT", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_paragraph("SMARTGARBAGE CHINTALAVALASA", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, color=(0, 100, 0))
add_paragraph("Community-Based Smart Waste Management\nand Digital Governance System", size=14, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

add_paragraph("Submitted by", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

names = [
    ("Name 1", "Register Number 1"),
    ("Name 2", "Register Number 2"),
    ("Name 3", "Register Number 3"),
    ("Name 4", "Register Number 4"),
]
for name, reg in names:
    add_paragraph(f"{name}  ({reg})", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

doc.add_paragraph()
add_paragraph("In partial fulfillment for the award of the degree of", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_paragraph("BACHELOR OF TECHNOLOGY", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_paragraph("IN", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_paragraph("COMPUTER SCIENCE & ENGINEERING", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_paragraph("(Artificial Intelligence & Machine Learning)", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_paragraph("Under the esteemed Guidance of", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_paragraph("GUIDE NAME", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_paragraph("DESIGNATION", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

doc.add_paragraph()
add_paragraph("DEPARTMENT OF DATA ENGINEERING", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_paragraph("MAHARAJ VIJAYARAM GAJAPATHI RAJ COLLEGE OF ENGINEERING (Autonomous)", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
add_paragraph("(Approved by AICTE, New Delhi, and permanently affiliated to JNTUGV, Vizianagaram),\nListed u/s 2(f) & 12(B) of UGC Act 1956.", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_paragraph("Vijayaram Nagar Campus, Chintalavalasa, Vizianagaram-535005, Andhra Pradesh", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_paragraph("October, 2025", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ════════════════════════════════════════════════════════════════════
# CERTIFICATE
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("CERTIFICATE")
doc.add_paragraph()
add_body('This is to certify that the project entitled "SmartGarbage Chintalavalasa — Community-Based Smart Waste Management and Digital Governance System" is the bonafide work carried out by Name 1 (Register Number 1), Name 2 (Register Number 2), Name 3 (Register Number 3), and Name 4 (Register Number 4), of B.Tech V Sem CSE-AIML, M.V.G.R. College of Engineering (Autonomous), Vizianagaram, during the year 2025-2026, in partial fulfilment of the requirements for the award of the Degree of Bachelor of Technology and that the project has not formed the basis for the award previously of any degree or any other similar title.')
doc.add_paragraph()
doc.add_paragraph()
add_paragraph("Signature of Project Guide", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3)
add_paragraph("Name\nDesignation\nDepartment: Data Engineering", size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=24)
doc.add_paragraph()
add_paragraph("Signature of Head of the Department", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3)
add_paragraph("Name\nDesignation\nDepartment: Data Engineering", size=12, align=WD_ALIGN_PARAGRAPH.LEFT)

# ════════════════════════════════════════════════════════════════════
# DECLARATION
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("DECLARATION")
doc.add_paragraph()
add_body('We hereby declare that the work done on the dissertation entitled "SmartGarbage Chintalavalasa — Community-Based Smart Waste Management and Digital Governance System" has been carried out by us and submitted in partial fulfilment for the award of credits in Bachelor of Technology in Computer Science and Engineering (Artificial Intelligence & Machine Learning) of M.V.G.R College of Engineering (Autonomous) and affiliated to JNTUGV, Vizianagaram. The various contents incorporated in the dissertation have not been submitted for the award of any degree of any other institution or university.')
doc.add_paragraph()
for name, reg in names:
    add_paragraph(f"{name} ({reg})", size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=12)

# ════════════════════════════════════════════════════════════════════
# ACKNOWLEDGEMENT
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("ACKNOWLEDGEMENT")
doc.add_paragraph()
add_body("We express our sincere gratitude to our project guide for their invaluable guidance and support as our mentor throughout the project. Their unwavering commitment to excellence and constructive feedback motivated us to achieve our project goals. We are greatly indebted to them for their exceptional guidance.")
add_body("Additionally, we extend our thanks to Prof. P.S. Sitharama Raju (Director), Dr. Y.M.C. Shekar (Principal), and Dr. Jyothi (Head of the Department) for their unwavering support and assistance, which were instrumental in the successful completion of the project. We are thankful for and fortunate enough to get constant encouragement and guidance from our Project Coordinator.")
add_body("We also acknowledge the dedicated assistance provided by all the staff members in the Department of Data Engineering. Finally, we appreciate the contributions of all those who directly or indirectly contributed to the successful execution of this endeavor.")
doc.add_paragraph()
for name, reg in names:
    add_paragraph(f"{name} ({reg})", size=12, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)

# ════════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("ABSTRACT")
doc.add_paragraph()
add_body("Rapid urbanization and population growth in semi-urban regions such as Chintalavalasa, Andhra Pradesh, have intensified waste management challenges. Existing systems rely heavily on manual collection schedules, paper-based complaint tracking, and limited citizen participation, leading to overflowing bins, delayed complaint resolution, and environmental degradation.")
add_body("This project presents SmartGarbage Chintalavalasa, a community-based smart waste management and digital governance system designed for the Chintalavalasa Gram Panchayat. The system integrates multiple technologies — a Flask-based web application, PostgreSQL database (hosted on Supabase), Internet of Things (IoT) smart-bin telemetry, machine learning (ML) predictions using a RandomForest regressor, Progressive Web App (PWA) offline capabilities, and a pay-as-you-throw (PAYT) billing mechanism — into a single, unified platform accessible to citizens, administrators, and sanitation workers.")
add_body("The platform provides five core portals: a Public Portal for waste collection schedules and complaint filing; a Citizen Portal for tracking complaints and earning Green Points for proper waste segregation; an Admin Portal for real-time dashboard monitoring, worker dispatch, and analytics; and a Worker Portal for GPS-tracked field operations. IoT-enabled smart bins report fill levels, battery status, temperature, and methane readings through authenticated API endpoints. A machine learning module predicts overflow risk based on historical and synthetic data, enabling proactive dispatch. The system supports bilingual communication (English and Telugu), web push notifications for complaint status updates, background job processing via Redis Queue, and comprehensive security measures including OWASP-aligned protections.")
add_body("Testing and evaluation indicate that the prototype achieves a homepage Time to First Byte (TTFB) of approximately 0.5 seconds, serves compressed responses via Brotli/Gzip, implements all nine recommended HTTP security headers, and passes automated accessibility checks. While the current deployment uses synthetic data for ML model training and simulated IoT telemetry, the architecture is designed for seamless integration with real-world sensor data and community-scale deployment. The project demonstrates how low-cost, open-source digital infrastructure can improve waste management governance in semi-urban Indian communities.")

add_paragraph("Keywords: Smart Waste Management, IoT, Machine Learning, PWA, Flask, Community Governance, PAYT Billing, Digital India", size=12, italic=True, space_after=12)

# ════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("TABLE OF CONTENTS")
doc.add_paragraph()

toc_items = [
    ("List of Abbreviations", ""),
    ("List of Figures", ""),
    ("List of Tables", ""),
    ("1. Introduction", ""),
    ("    1.1 Problem Statement", ""),
    ("    1.2 Project Objective", ""),
    ("    1.3 Scope of the Project", ""),
    ("2. Literature Survey", ""),
    ("3. Data Gathering / Data Used", ""),
    ("4. Methodology / System Design", ""),
    ("5. Implementation / Modules", ""),
    ("6. Results / Outputs", ""),
    ("7. Impact Assessment", ""),
    ("8. Challenges Faced", ""),
    ("9. Conclusion", ""),
    ("10. Future Work", ""),
    ("References", ""),
    ("Appendix A: Packages, Tools Used & Working Process", ""),
    ("Appendix B: Source Code", ""),
]
for title, pg in toc_items:
    add_paragraph(title, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

# ════════════════════════════════════════════════════════════════════
# LIST OF ABBREVIATIONS
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("LIST OF ABBREVIATIONS")
doc.add_paragraph()
abbrevs = [
    ("AI", "Artificial Intelligence"),
    ("API", "Application Programming Interface"),
    ("ARIA", "Accessible Rich Internet Applications"),
    ("BLoC", "Business Logic Component"),
    ("CDN", "Content Delivery Network"),
    ("CSP", "Content Security Policy"),
    ("CSS", "Cascading Style Sheets"),
    ("DBMS", "Database Management System"),
    ("ETL", "Extract, Transform, Load"),
    ("FCP", "First Contentful Paint"),
    ("GDPR", "General Data Protection Regulation"),
    ("GPS", "Global Positioning System"),
    ("Gunicorn", "Green Unicorn WSGI HTTP Server"),
    ("HSTS", "HTTP Strict Transport Security"),
    ("HTML", "HyperText Markup Language"),
    ("HTTP", "HyperText Transfer Protocol"),
    ("IoT", "Internet of Things"),
    ("JSON-LD", "JSON for Linked Data"),
    ("JSA", "JavaScript Application"),
    ("ML", "Machine Learning"),
    ("MFA", "Multi-Factor Authentication"),
    ("MVC", "Model-View-Controller"),
    ("OAuth", "Open Authorization"),
    ("OTP", "One-Time Password"),
    ("OWASP", "Open Web Application Security Project"),
    ("PAYT", "Pay-As-You-Throw"),
    ("PWA", "Progressive Web App"),
    ("RBAC", "Role-Based Access Control"),
    ("REST", "Representational State Transfer"),
    ("RQ", "Redis Queue"),
    ("SBM", "Swachh Bharat Mission"),
    ("SEO", "Search Engine Optimization"),
    ("SQL", "Structured Query Language"),
    ("SSL", "Secure Sockets Layer"),
    ("SW", "Service Worker"),
    ("TTFB", "Time to First Byte"),
    ("URI", "Uniform Resource Identifier"),
    ("VAPID", "Voluntary Application Server Identification"),
    ("WCAG", "Web Content Accessibility Guidelines"),
    ("WCS", "Web Content Security"),
    ("WGS84", "World Geodetic System 1984"),
    ("XML", "Extensible Markup Language"),
]
add_table(["Abbreviation", "Full Form"], abbrevs, col_widths=[1.5, 5.0])

# ════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("LIST OF FIGURES")
doc.add_paragraph()
figures = [
    ("Figure 1.1", "Problem-to-Solution Overview", ""),
    ("Figure 4.1", "Development Methodology Flowchart", ""),
    ("Figure 4.2", "Overall System Architecture", ""),
    ("Figure 4.3", "System Workflow", ""),
    ("Figure 4.4", "Complaint Lifecycle Flowchart", ""),
    ("Figure 4.5", "Data Flow Diagram", ""),
    ("Figure 4.6", "Database ER Diagram", ""),
    ("Figure 4.7", "Machine Learning Pipeline", ""),
    ("Figure 4.8", "IoT Data Flow Diagram", ""),
    ("Figure 4.9", "PWA Offline Workflow", ""),
    ("Figure 6.1", "Homepage Output", ""),
    ("Figure 6.2", "Collection Schedule Output", ""),
    ("Figure 6.3", "Complaint Reporting Output", ""),
    ("Figure 6.4", "Citizen Dashboard Output", ""),
    ("Figure 6.5", "Admin Dashboard Output", ""),
    ("Figure 6.6", "Worker Dispatch Output", ""),
    ("Figure 6.7", "IoT Telemetry Output", ""),
    ("Figure 6.8", "ML Prediction Output", ""),
]
add_table(["Figure No.", "Title", "Page No."], figures, col_widths=[1.2, 4.0, 1.0])

# ════════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("LIST OF TABLES")
doc.add_paragraph()
tables_list = [
    ("Table 2.1", "Literature Survey Summary", ""),
    ("Table 3.1", "Data Sources and Collection Methods", ""),
    ("Table 3.2", "Ward Information", ""),
    ("Table 3.3", "Database Schema Overview", ""),
    ("Table 4.1", "Technology Stack", ""),
    ("Table 4.2", "Stakeholder Requirements", ""),
    ("Table 5.1", "Module Summary", ""),
    ("Table 6.1", "Implemented Features Summary", ""),
    ("Table 6.2", "Performance Evaluation Results", ""),
    ("Table 6.3", "Security Header Evaluation", ""),
    ("Table 7.1", "Impact Assessment Metrics", ""),
    ("Table 8.1", "Challenges and Solutions", ""),
]
add_table(["Table No.", "Title", "Page No."], tables_list, col_widths=[1.2, 4.0, 1.0])


# ════════════════════════════════════════════════════════════════════
# CHAPTER 1 — INTRODUCTION
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("1. INTRODUCTION")

add_heading_section("1.1 Problem Statement")
add_body("Chintalavalasa is a semi-urban panchayat in Vizianagaram district, Andhra Pradesh, with a population of approximately 12,000 residents spread across five administrative wards. The current waste management system relies on manual collection schedules that are communicated verbally or through informal channels, resulting in inconsistent service delivery. Citizens have no reliable mechanism to report missed collections or overflowing bins, and complaint resolution is tracked informally without transparency or accountability.")
add_body("Key problems identified through community interaction and field observation include: (a) lack of centralized collection schedules accessible to all residents; (b) absence of a digital grievance redressal mechanism for waste-related complaints; (c) no real-time monitoring of bin fill levels, leading to overflow situations; (d) no data-driven approach to predict collection needs or allocate resources efficiently; (e) limited citizen engagement in waste segregation and recycling; and (f) absence of financial accountability through usage-based billing.")

add_heading_section("1.2 Project Objective")
add_body("The primary objective of this project is to design, develop, and deploy a community-based smart waste management and digital governance system for Chintalavalasa Gram Panchayat. The specific objectives are:")
objectives = [
    "To develop a web-based platform providing waste collection schedules, complaint filing, and tracking for citizens.",
    "To implement an admin dashboard for real-time monitoring of complaints, worker dispatch, and analytics.",
    "To integrate IoT smart-bin telemetry for monitoring fill levels, battery status, and environmental conditions.",
    "To develop a machine learning module for predicting bin overflow risk and enabling proactive dispatch.",
    "To implement a pay-as-you-throw (PAYT) billing mechanism to incentivize waste segregation.",
    "To provide Progressive Web App (PWA) capabilities with offline support for areas with limited connectivity.",
    "To implement comprehensive security measures aligned with OWASP recommendations.",
    "To support bilingual communication in English and Telugu for wider accessibility.",
]
for obj in objectives:
    add_paragraph(f"\u2022  {obj}", size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

add_heading_section("1.3 Scope of the Project")
add_body("The SmartGarbage system encompasses the following scope:")
add_body("\u2022  In Scope: Web application with public, citizen, admin, and worker portals; IoT smart-bin telemetry integration; machine learning overflow prediction; PAYT billing with UPI/Razorpay integration; Green Points gamification; PWA offline support; push notifications; bilingual support; and comprehensive security architecture.")
add_body("\u2022  Out of Scope: Native mobile applications (iOS/Android); physical IoT hardware manufacturing; integration with external municipal corporation APIs; blockchain-based carbon credit systems; WhatsApp/SMS gateway production deployment (integrated but requiring third-party API keys); and real-world IoT sensor deployment at community scale.")
add_body("The system is designed as a prototype demonstration for the Chintalavalasa community, with architecture that supports scaling to additional panchayats and integration with real sensor hardware in future phases.")


# ════════════════════════════════════════════════════════════════════
# CHAPTER 2 — LITERATURE SURVEY
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("2. LITERATURE SURVEY")

add_body("A comprehensive review of existing literature and systems was conducted to understand the current state of waste management solutions and identify gaps that SmartGarbage addresses.")

add_heading_section("2.1 Existing Waste Management Approaches")
add_body("Traditional waste management in Indian semi-urban areas relies on manual collection with fixed schedules, paper-based complaint registers, and verbal communication between citizens and sanitation workers. The Swachh Bharat Mission (SBM) Grameen Phase II has promoted source segregation, pay-as-you-throw models, and digital monitoring in rural areas, but implementation remains inconsistent at the panchayat level. Existing approaches suffer from lack of transparency, delayed grievance resolution, and absence of data-driven decision making.")

add_heading_section("2.2 Digital Waste Management Systems")
add_body("Several digital platforms have been developed for waste management. SBM Urban (India) provides a complaint registration and tracking system for urban areas, but lacks IoT integration, ML prediction, and offline capabilities. VA.gov (USA) demonstrates task-based digital service design with strong accessibility compliance but does not address waste management specifically. GOV.UK (UK) sets the benchmark for government digital services with its Service Standard, emphasizing user needs, accessibility (WCAG 2.1), and performance optimization. SmartGarbage draws design inspiration from these platforms while adding domain-specific features for waste management.")

add_heading_section("2.3 IoT-Based Smart Waste Management")
add_body("Research by Anagnostopoulos et al. (2017) demonstrated that IoT-based waste monitoring using ultrasonic fill-level sensors can reduce collection costs by 30-50% through optimized routing. Md Azree ad et al. (2020) proposed a smart bin system with real-time monitoring and automated collection scheduling. Kumar and Sharma (2021) reviewed IoT applications in solid waste management and identified that most solutions focus on individual components (sensor, routing, or prediction) rather than integrated platforms. SmartGarbage addresses this gap by integrating IoT telemetry with complaint management, citizen engagement, and ML prediction in a unified system.")

add_heading_section("2.4 Machine Learning for Waste Prediction")
add_body("Afshin et al. (2021) applied gradient boosting regression to predict municipal solid waste generation using demographic and seasonal features. Chen et al. (2020) demonstrated that Random Forest and Gradient Boosting models achieve comparable accuracy for short-term waste volume prediction. Thung and Yang (2016) reviewed waste classification using deep learning for automated sorting. SmartGarbage implements a RandomForest regressor for bin overflow prediction, using ward-level features including day-of-week, season, recent complaint volume, and historical patterns, with transparent fallback heuristics when the model is unavailable.")

add_heading_section("2.5 Accessibility and Government Digital Standards")
add_body("The Web Content Accessibility Guidelines (WCAG) 2.1 Level AA provide the international standard for web accessibility. The GOV.UK Service Standard mandates that digital services must be accessible to all users, including those with disabilities. OWASP Top 10 (2021) identifies the most critical web application security risks. SmartGarbage implements ARIA landmarks, skip-to-content links, keyboard navigation, sufficient color contrast, and semantic HTML to meet accessibility requirements, while applying OWASP-recommended security headers including Content Security Policy, HSTS, and X-Content-Type-Options.")

add_heading_section("2.6 Research / Implementation Gap")
add_body("Existing solutions address individual aspects of waste management: IoT monitoring systems, citizen complaint portals, or ML prediction models. However, there is no integrated, low-cost platform that combines: (a) citizen grievance reporting with GPS and photo evidence; (b) collection schedule management; (c) IoT smart-bin telemetry; (d) ML-based overflow prediction; (e) PAYT billing with gamification; (f) PWA offline support for low-connectivity areas; and (g) comprehensive security and accessibility — all in a single platform designed for semi-urban Indian communities. SmartGarbage fills this gap.")

add_heading_section("2.7 Proposed Contribution")
add_body("SmartGarbage contributes an integrated, open-source waste management platform that: combines citizen engagement, administrative oversight, worker coordination, IoT monitoring, and ML prediction; provides offline-first PWA capabilities for low-connectivity areas; implements PAYT billing with Green Points gamification to incentivize segregation; supports bilingual communication; and achieves strong security and accessibility compliance — all within a deployable, low-cost architecture suitable for semi-urban Indian panchayats.")


# ════════════════════════════════════════════════════════════════════
# CHAPTER 3 — DATA GATHERING / DATA USED
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("3. DATA GATHERING / DATA USED")

add_heading_section("3.1 Study Area / Community Profile")
add_body("Chintalavalasa is a semi-urban panchayat located in Vizianagaram district, Andhra Pradesh, India. The panchayat serves approximately 12,000 residents across five administrative wards. The community includes residential areas, a college zone (MVGR College of Engineering), commercial areas along the junction, and mixed-use zones. The geographic coordinates range from latitude 18.0552 to 18.0751 and longitude 83.4005 to 83.4201.")

add_heading_section("3.2 Data Collection Methods")
add_body("Data for this project was collected through the following methods:")
add_body("\u2022  Community Interaction: Direct conversations with residents, ward members, and sanitation workers to understand waste management pain points and requirements.")
add_body("\u2022  Field Observation: On-site observation of waste collection routes, bin locations, and current complaint handling processes.")
add_body("\u2022  Administrative Data: Ward boundaries, population estimates, and existing infrastructure information from the Gram Panchayat office.")
add_body("\u2022  System-Generated Data: Test data generated during development for ML model training, including synthetic IoT telemetry, complaint records, and waste declaration entries.")
add_body("\u2022  Public Records: Government waste management guidelines from SBM Grameen Phase II, existing digital platform documentation (GOV.UK, SBM Urban), and accessibility standards (WCAG 2.1).")

add_heading_section("3.3 Data Sources")
data_sources = [
    ("Community Surveys", "Qualitative", "Resident and worker interviews", "Requirements gathering"),
    ("Field Observations", "Qualitative", "On-site waste collection observation", "System design"),
    ("Administrative Records", "Semi-structured", "Ward boundaries, population data", "Study area profile"),
    ("Synthetic ML Data", "Quantitative", "600-row synthetic grid", "ML model training"),
    ("Synthetic IoT Data", "Quantitative", "Simulated sensor readings", "IoT module testing"),
    ("Test User Data", "Quantitative", "Registration and complaint records", "System testing"),
]
add_table(["Data Source", "Type", "Description", "Use in Project"], data_sources, col_widths=[1.5, 1.2, 2.2, 1.5])

add_heading_section("3.4 Ward Information")
ward_data = [
    ("Ward 1", "MVGR College Area", "18.0552", "83.4051"),
    ("Ward 2", "Chintalavalasa Junction", "18.0675", "83.4094"),
    ("Ward 3", "RTC Colony", "18.0702", "83.4153"),
    ("Ward 4", "Ramalayam Street", "18.0650", "83.4005"),
    ("Ward 5", "Sai Nagar", "18.0751", "83.4201"),
]
add_table(["Ward ID", "Ward Name", "Latitude", "Longitude"], ward_data, col_widths=[1.0, 2.5, 1.5, 1.5])

add_heading_section("3.5 Data Used by the Application")
add_body("The SmartGarbage system manages six categories of operational data:")
add_body("\u2022  User Data: Citizen registrations, roles (citizen/worker/admin), OTP verification records, and Green Points balances.")
add_body("\u2022  Complaint Data: Citizen-filed complaints with GPS coordinates, photo attachments, status history, and resolution timestamps.")
add_body("\u2022  Schedule Data: Ward-specific waste collection schedules with day, time, and collection type information.")
add_body("\u2022  IoT Telemetry: Smart-bin readings including fill level (0-100%), battery level, temperature, methane concentration, and overflow ETA.")
add_body("\u2022  Waste Declaration Data: Citizen-reported waste quantities categorized as wet (organic), dry (recyclable), sanitary, and hazardous.")
add_body("\u2022  Billing Data: PAYT invoices with period, weight, compliance score, penalty multiplier, and payment status.")

add_heading_section("3.6 Data Preparation")
add_body("For the machine learning module, a synthetic training dataset of 600 rows was generated because sufficient historical waste telemetry was unavailable during prototype development. The synthetic data was created to demonstrate the prediction pipeline and integration architecture. Features include day-of-week, season index, recent complaint volume in the ward, ward identifier, and historical fill patterns. The model is designed to accept real historical data when available, at which point the synthetic dataset would be replaced through the standard ETL pipeline.")

add_heading_section("3.7 Database Design")
add_body("The system uses PostgreSQL (hosted on Supabase) with SQLAlchemy ORM. The database schema comprises the following primary entities:")
schema_data = [
    ("User", "id, username, email, password_hash, role, phone, green_points, otp, is_approved, email_verified"),
    ("Complaint", "id, name, phone, ward, address, description, photo, status, latitude, longitude, user_id, created_at"),
    ("ComplaintStatusLog", "id, complaint_id, status, note, created_at"),
    ("SmartBin", "id, hardware_id, latitude, longitude, level, battery_level, temperature, methane, status, ward"),
    ("WorkerProfile", "id, user_id, vehicle_id, latitude, longitude, status, performance_rating"),
    ("Schedule", "id, ward, day, time, collection_type"),
    ("WasteDeclaration", "id, user_id, wet_kg, dry_kg, sanitary_kg, hazardous_kg, ward, timestamp"),
    ("PAYTInvoice", "id, user_id, period, weight_kg, amount_rs, status, issued_at, paid_at"),
    ("PushSubscription", "id, user_id, endpoint, p256dh, auth"),
    ("NotificationPreference", "id, user_id, complaint_submitted, complaint_assigned, etc."),
]
add_table(["Entity", "Key Attributes"], schema_data, col_widths=[1.8, 4.7])


# ════════════════════════════════════════════════════════════════════
# CHAPTER 4 — METHODOLOGY / SYSTEM DESIGN
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("4. METHODOLOGY / SYSTEM DESIGN")

add_heading_section("4.1 Requirement Analysis")
add_body("Requirements were gathered through community interaction, field observation, and analysis of existing government waste management frameworks (SBM Grameen Phase II). The requirements were categorized into functional requirements (FR) and non-functional requirements (NFR):")

add_heading_subsection("Functional Requirements")
fr_items = [
    ("FR-01", "Public Portal", "Display collection schedules, allow anonymous complaint filing"),
    ("FR-02", "Citizen Portal", "Track complaints, manage profile, earn Green Points"),
    ("FR-03", "Admin Portal", "Dashboard with analytics, worker dispatch, complaint management"),
    ("FR-04", "Worker Portal", "GPS tracking, dispatch acceptance, photo evidence upload"),
    ("FR-05", "IoT Integration", "Smart-bin telemetry ingestion and monitoring"),
    ("FR-06", "ML Prediction", "Overflow risk prediction with transparent fallback"),
    ("FR-07", "PAYT Billing", "Usage-based invoicing with UPI/Razorpay payment"),
    ("FR-08", "Push Notifications", "Web push alerts for complaint status changes"),
    ("FR-09", "Offline Support", "PWA with IndexedDB offline queue"),
    ("FR-10", "Bilingual Support", "English and Telugu language options"),
]
add_table(["ID", "Module", "Description"], fr_items, col_widths=[0.8, 1.5, 4.2])

add_heading_subsection("Non-Functional Requirements")
nfr_items = [
    ("NFR-01", "Performance", "TTFB < 1 second, compressed responses"),
    ("NFR-02", "Security", "OWASP-aligned headers, RBAC, bcrypt, OTP"),
    ("NFR-03", "Accessibility", "WCAG 2.1 AA compliance, ARIA landmarks"),
    ("NFR-04", "Scalability", "Horizontal scaling via gunicorn workers"),
    ("NFR-05", "Offline", "PWA with service worker and IndexedDB"),
]
add_table(["ID", "Category", "Description"], nfr_items, col_widths=[0.8, 1.5, 4.2])

add_heading_section("4.2 System Architecture")
add_body("The system follows a layered architecture pattern with clear separation of concerns:")

add_figure_caption("Figure 4.2: Overall System Architecture")
arch_text = (
    "The architecture comprises the following layers:\n\n"
    "1. Presentation Layer: Browser/PWA interface using Jinja2 templates with Bootstrap 5 and custom CSS.\n"
    "2. Application Layer: Flask application with route modules for public, citizen, admin, and worker portals.\n"
    "3. Business Logic Layer: ML prediction engine, background job processor (Redis Queue), push notification service, and PAYT billing calculator.\n"
    "4. Data Layer: PostgreSQL database via SQLAlchemy ORM with connection pooling and caching (Redis + in-process).\n"
    "5. External Services: Open-Meteo weather API, Supabase hosting, Sentry error tracking, and optional Twilio/WhatsApp integration.\n\n"
    "Each layer communicates through well-defined interfaces. The Flask application handles HTTP requests and delegates to the appropriate route module. Route modules interact with the data layer through SQLAlchemy models and call business logic services for ML predictions, notifications, and billing. Background jobs run asynchronously via Redis Queue when configured."
)
add_body(arch_text)

add_heading_section("4.3 Development Methodology")
add_body("The project followed an iterative development approach with the following phases:")
add_body("\u2022  Phase 1 — Requirements Gathering (Weeks 1-2): Community interaction, field observation, and requirements documentation.")
add_body("\u2022  Phase 2 — System Design (Weeks 3-4): Architecture design, database schema, UI wireframes, and technology selection.")
add_body("\u2022  Phase 3 — Core Development (Weeks 5-10): Implementation of public portal, citizen portal, admin portal, and worker portal.")
add_body("\u2022  Phase 4 — Integration (Weeks 11-14): IoT telemetry integration, ML model training and deployment, PAYT billing, and background jobs.")
add_body("\u2022  Phase 5 — Enhancement (Weeks 15-18): PWA offline support, push notifications, bilingual support, security hardening, and accessibility.")
add_body("\u2022  Phase 6 — Testing and Deployment (Weeks 19-20): Functional testing, performance testing, security evaluation, and production deployment.")

add_heading_section("4.4 Technology Stack")
tech_data = [
    ("Backend Framework", "Flask 3.1.3", "Python web framework"),
    ("Database", "PostgreSQL (Supabase)", "Relational database with connection pooling"),
    ("ORM", "SQLAlchemy 2.0.50", "Python SQL toolkit and ORM"),
    ("Frontend", "Bootstrap 5 + Custom CSS", "Responsive UI framework"),
    ("JavaScript", "Vanilla JS + Socket.IO", "Client-side interactivity"),
    ("ML Library", "scikit-learn 1.9.0", "RandomForest, preprocessing"),
    ("Data Processing", "pandas, numpy", "Data manipulation"),
    ("Task Queue", "Redis + RQ 2.2.0", "Background job processing"),
    ("WSGI Server", "Gunicorn 26.0.0", "Production HTTP server"),
    ("Push Notifications", "pywebpush 2.0.0", "Web Push API with VAPID"),
    ("Error Tracking", "Sentry SDK 2.5.0", "Production error monitoring"),
    ("Caching", "Flask-Cache + Redis", "Multi-layer caching"),
    ("Security", "Flask-Talisman, Flask-Limiter", "Security headers, rate limiting"),
    ("Email", "Flask-Mailman", "Email notifications"),
    ("PDF Generation", "ReportLab 5.0.0", "PAYT invoice generation"),
    ("Container", "Docker", "Deployment containerization"),
    ("Hosting", "Render.com", "Cloud platform hosting"),
]
add_table(["Component", "Technology", "Purpose"], tech_data, col_widths=[1.8, 2.2, 2.5])

add_heading_section("4.5 Data Flow")
add_body("The system manages multiple data flows. The primary complaint lifecycle flow is as follows:")
add_body("Citizen submits complaint (GPS + photo + description) -> Validation (field completeness, duplicate detection) -> Complaint stored in database -> Admin notification -> Admin assigns to worker -> Worker receives dispatch -> Worker visits location -> Worker uploads photo evidence -> Admin verifies resolution -> Complaint status updated to 'Resolved' -> Citizen notified via push/email -> Green Points awarded.")

add_body("The IoT telemetry flow operates as follows:")
add_body("Smart bin sensors generate readings (fill level, battery, temperature, methane) -> IoT device transmits via authenticated API -> Telemetry processed and stored in SmartBin table -> Admin dashboard displays real-time status -> ML model uses latest readings for overflow prediction -> Priority dispatch queue updated.")

add_heading_section("4.6 Machine Learning Methodology")
add_body("The ML module implements a RandomForest regressor for predicting the hours until a smart bin reaches 90% fill level (overflow risk). The model training pipeline includes:")
add_body("\u2022  Data Preparation: A synthetic dataset of 600 rows was generated because sufficient historical telemetry was unavailable during prototype development.")
add_body("\u2022  Feature Engineering: Features include day_of_week (0-6), season_idx (0-3), recent_complaint_count (integer), and ward_id (categorical, encoded).")
add_body("\u2022  Model Training: RandomForestRegressor from scikit-learn with 100 estimators, trained on the synthetic grid data.")
add_body("\u2022  Prediction: For a given ward and current conditions, the model predicts hours_until_90pct_fill. Bins are ranked by urgency for dispatch prioritization.")
add_body("\u2022  Fallback: When the model artifact is unavailable, a transparent heuristic based on average fill rate and recent complaint patterns is used.")
add_body("Note: Synthetic data was used during prototype development because sufficient historical waste telemetry was not available. The model demonstration validates the prediction pipeline and integration architecture, not real-world predictive accuracy. Real-world historical data integration is proposed for future work.")

add_heading_section("4.7 Security Architecture")
add_body("The security architecture implements OWASP-recommended protections:")
add_body("\u2022  Authentication: Flask-Login with bcrypt password hashing and OTP/MFA verification.")
add_body("\u2022  Authorization: Role-Based Access Control (RBAC) with three roles: citizen, worker, and admin. Admin approval required for new accounts.")
add_body("\u2022  HTTP Security Headers: All nine recommended headers — Content-Security-Policy, Strict-Transport-Security, X-Content-Type-Options, X-Frame-Options, X-XSS-Protection, Referrer-Policy, Permissions-Policy, Cross-Origin-Opener-Policy, and Cross-Origin-Embedder-Policy.")
add_body("\u2022  Rate Limiting: Flask-Limiter with configurable storage backend (Redis in production, memory in development).")
add_body("\u2022  Input Validation: Server-side validation for all user inputs using Flask-WTF forms with CSRF protection.")
add_body("\u2022  SQL Injection Prevention: All database queries use SQLAlchemy parameterized queries.")
add_body("\u2022  Consent Management: Anonymized GDPR/DPDP-style consent capture with salted SHA-256 fingerprints.")

add_heading_section("4.8 PWA and Offline Methodology")
add_body("The PWA implementation provides offline capabilities through the following mechanism:")
add_body("\u2022  Service Worker: A custom service worker (sw.js) with versioned precache manifest intercepts fetch requests and serves cached responses when offline.")
add_body("\u2022  IndexedDB Queue: When a citizen files a complaint while offline, the request is stored in IndexedDB and queued for background synchronization.")
add_body("\u2022  Background Sync: When connectivity is restored, queued requests are automatically submitted to the server.")
add_body("\u2022  Manifest: A comprehensive web app manifest enables installation as a native-like app with shortcuts, screenshots, and standalone display mode.")
add_body("\u2022  Splash Screen: A branded loading animation displays while the PWA initializes on first install.")
add_body("\u2022  Install Banner: A responsive PWA install prompt appears on mobile after the second visit, with iOS-specific fallback instructions.")


# ════════════════════════════════════════════════════════════════════
# CHAPTER 5 — IMPLEMENTATION / MODULES
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("5. IMPLEMENTATION / MODULES")

modules = [
    ("5.1 Public Portal", "Public Portal",
     "Provides waste-management information accessible to all visitors without authentication.",
     ["Waste collection schedule display by ward",
      "Anonymous complaint filing with GPS capture",
      "Complaint tracking via signed token URL",
      "Ward transparency dashboard with statistics",
      "FAQ, About, Contact, Privacy Policy, Terms pages",
      "Site-wide search across all pages",
      "RSS feed and sitemap for search engines",
      "AI-friendly llms.txt for chatbot visibility"],
     "Flask + Jinja2 + Bootstrap 5 + Vanilla JS",
     "GET /, /schedule, /transparency, /track/<token>, /faq, /about, /search",
     "Schedule page, complaint form, ward transparency dashboard",
     "Fully implemented and verified"),

    ("5.2 Citizen Portal", "Citizen Portal",
     "Authenticated portal for registered citizens to manage complaints and engagement.",
     ["User registration with OTP verification",
      "Complaint filing with photo upload and GPS",
      "Complaint tracking with status timeline",
      "Dashboard with personal complaint history",
      "Green Points balance and leaderboard",
      "Waste declaration (segregation reporting)",
      "Notification preferences management",
      "Push notification subscription"],
     "Flask-Login + SQLAlchemy + Socket.IO",
     "POST /register, /report, GET /dashboard, /notifications/preferences",
     "Citizen dashboard, complaint form, notification preferences",
     "Fully implemented and verified"),

    ("5.3 Admin Portal", "Admin Portal",
     "Comprehensive dashboard for administrators to manage operations.",
     ["Real-time complaint overview with status filters",
      "Worker assignment and dispatch management",
      "IoT smart-bin monitoring dashboard",
      "ML prediction display with urgency ranking",
      "PAYT invoice management",
      "Push notification analytics",
      "Audit log viewer",
      "User management (approve/reject accounts)",
      "Bot simulator for testing notifications",
      "Performance KPI tracking"],
     "Flask-Login + RBAC + Socket.IO + Chart.js",
     "GET /admin, POST /admin/assign, /admin/dispatch",
     "Admin control room, analytics dashboard, push notification section",
     "Fully implemented and verified"),

    ("5.4 Worker Portal", "Worker Portal",
     "Mobile-friendly portal for sanitation workers.",
     ["Dispatch acceptance and status updates",
      "GPS location tracking during field operations",
      "Photo evidence upload for completed work",
      "Task queue with priority ordering",
      "Worker profile and performance rating"],
     "Flask-Login + Geolocation API",
     "GET /worker, POST /worker/update-status",
     "Worker dashboard, dispatch view",
     "Fully implemented and verified"),

    ("5.5 IoT Smart Bin Module", "IoT Smart Bin Module",
     "Handles IoT telemetry ingestion and smart-bin monitoring.",
     ["Authenticated API for sensor data submission",
      "Fill level monitoring (0-100%)",
      "Battery level and temperature tracking",
      "Methane concentration monitoring",
      "Automated status classification (Safe/Warning/Critical)",
      "Sensor fault detection",
      "Waste stream categorization"],
     "Flask API + SQLAlchemy + Structured Logging",
     "POST /api/iot/telemetry, GET /admin (IoT section)",
     "IoT telemetry display on admin dashboard",
     "Implemented with simulated data"),

    ("5.6 Machine Learning Module", "Machine Learning Module",
     "Predicts bin overflow risk using ML regression.",
     ["RandomForestRegressor for overflow prediction",
      "Feature engineering (day, season, complaints, ward)",
      "Transparent heuristic fallback",
      "Model persistence via pickle",
      "Integration with dispatch prioritization"],
     "scikit-learn + pandas + numpy",
     "predict_miss(ward), /schedule (POST)",
     "ML prediction display in admin and schedule pages",
     "Implemented with synthetic training data"),

    ("5.7 Green Points Module", "Green Points Module",
     "Gamification system incentivizing waste segregation.",
     ["Points earned for waste declarations",
      "Streak tracking for consecutive segregated declarations",
      "Leaderboard display",
      "Redemption mechanism"],
     "Flask-Login + SQLAlchemy",
     "POST /waste-declare, GET /dashboard (Green Points section)",
     "Green Points display on citizen dashboard",
     "Fully implemented"),

    ("5.8 PAYT Module", "Pay-As-You-Throw Module",
     "Usage-based billing incentivizing proper waste segregation.",
     ["Invoice generation based on waste weight",
      "Compliance scoring and penalty calculation",
      "UPI and Razorpay payment integration",
      "PDF invoice generation via ReportLab",
      "Dunning reminder automation"],
     "Flask + ReportLab + Razorpay SDK",
     "POST /payt/invoice, GET /payt/pay/<id>",
     "PAYT invoice page, payment integration",
     "Implemented with test payment mode"),

    ("5.9 Background Jobs Module", "Background Jobs Module",
     "Asynchronous task processing via Redis Queue.",
     ["Complaint status change notifications",
      "SLA escalation alerts",
      "Failed complaint sweep",
      "PAYT dunning reminders",
      "Email delivery",
      "Push notification dispatch"],
     "Redis + RQ + Background Workers",
     "schedule_dunning(), schedule_sla_escalation()",
     "Background job execution in production",
     "Fully implemented (no-op without Redis)"),

    ("5.10 PWA and Offline Module", "PWA and Offline Module",
     "Progressive Web App capabilities for offline access.",
     ["Service worker with precache manifest",
      "IndexedDB offline complaint queue",
      "Background sync on reconnection",
      "Web app manifest with shortcuts",
      "Splash screen on PWA install",
      "Install prompt banner for mobile",
      "Standalone display mode"],
     "Service Worker API + IndexedDB + Web Manifest",
     "sw.js, manifest.json, /offline",
     "PWA installation and offline queue",
     "Fully implemented and verified"),
]

for heading, title, purpose, functions, tech, routes, evidence, status in modules:
    add_heading_section(heading)
    add_heading_subsection("Purpose")
    add_body(purpose)
    add_heading_subsection("Functions")
    for f in functions:
        add_paragraph(f"\u2022  {f}", size=12, space_after=3)
    add_heading_subsection("Technology")
    add_body(tech)
    add_heading_subsection("Key Routes")
    add_body(routes)
    add_heading_subsection("Evidence")
    add_body(evidence)
    add_heading_subsection("Implementation Status")
    add_body(status)


# ════════════════════════════════════════════════════════════════════
# CHAPTER 6 — RESULTS / OUTPUTS
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("6. RESULTS / OUTPUTS")

add_heading_section("6.1 Implemented Features Summary")
features_summary = [
    ("Public Portal", "Homepage, schedule, transparency, search, FAQ", "Fully Implemented", "200 OK"),
    ("Citizen Portal", "Registration, complaint filing, dashboard, Green Points", "Fully Implemented", "200 OK"),
    ("Admin Portal", "Dashboard, dispatch, IoT monitoring, ML display", "Fully Implemented", "200 OK"),
    ("Worker Portal", "Dispatch, GPS tracking, evidence upload", "Fully Implemented", "200 OK"),
    ("IoT Integration", "Telemetry API, sensor monitoring", "Implemented (Simulated)", "API functional"),
    ("ML Prediction", "RandomForest overflow prediction", "Implemented (Synthetic Data)", "Pipeline functional"),
    ("PAYT Billing", "Invoice generation, payment integration", "Implemented (Test Mode)", "Invoice generation OK"),
    ("Push Notifications", "Web push with VAPID, preferences, analytics", "Fully Implemented", "API functional"),
    ("PWA / Offline", "Service worker, IndexedDB queue, manifest", "Fully Implemented", "All endpoints 200"),
    ("Security", "OWASP headers, RBAC, bcrypt, OTP", "Fully Implemented", "9/9 headers present"),
    ("Accessibility", "ARIA landmarks, skip-to-content, contrast", "Implemented", "Automated checks pass"),
    ("Bilingual Support", "English and Telugu", "Implemented", "Language toggle functional"),
]
add_table(["Module", "Key Features", "Status", "Verification"], features_summary, col_widths=[1.3, 2.5, 1.5, 1.2])

add_heading_section("6.2 Homepage Output")
add_body("The homepage displays the collection schedule lookup, complaint filing shortcut, ward transparency map, weather widget, community impact statistics (total wards, active smart bins, resolved complaints, average resolution time), and a public search bar. The homepage loads in approximately 0.5 seconds with Brotli compression and caching headers.")
add_figure_caption("Figure 6.1: Homepage Output — Collection schedule, community impact, and weather widget")

add_heading_section("6.3 Collection Schedule Output")
add_body("The schedule page allows citizens to select their ward from a dropdown and view the collection schedule for that ward. When a ward is selected, the system also runs the ML prediction to display the estimated overflow risk. The schedule shows day, time, and collection type for each entry.")
add_figure_caption("Figure 6.2: Collection Schedule Output — Ward selection with ML prediction display")

add_heading_section("6.4 Complaint Reporting Output")
add_body("The complaint form captures the reporter's name, phone, ward selection, address, description, optional photo attachment, and automatic GPS coordinates. On submission, the complaint is stored with status 'Submitted' and a signed tracking token is generated for anonymous tracking.")
add_figure_caption("Figure 6.3: Complaint Reporting Output — GPS-enabled complaint form with photo upload")

add_heading_section("6.5 Complaint Tracking Output")
add_body("The tracking page (accessible via the signed token URL) displays the complaint status timeline showing each status transition with timestamp and notes. It also shows the ward's average resolution time for transparency.")
add_figure_caption("Figure 6.4: Complaint Tracking Output — Status timeline with SLA information")

add_heading_section("6.6 Citizen Dashboard Output")
add_body("The citizen dashboard shows the user's complaint history, Green Points balance, recent waste declarations, and notification preferences. It provides quick links to file new complaints and view schedules.")
add_figure_caption("Figure 6.5: Citizen Dashboard Output — Complaint history, Green Points, and quick actions")

add_heading_section("6.7 Admin Dashboard Output")
add_body("The admin control room displays real-time complaint statistics, IoT smart-bin status with fill levels and battery readings, worker assignment queue, and push notification analytics. The dashboard uses Socket.IO for live updates.")
add_figure_caption("Figure 6.6: Admin Dashboard Output — Real-time monitoring with IoT and push analytics")

add_heading_section("6.8 Worker Dispatch Output")
add_body("The worker portal shows assigned tasks with priority ranking, GPS tracking status, and photo evidence upload functionality. Workers can accept or complete dispatches directly from their mobile device.")
add_figure_caption("Figure 6.7: Worker Dispatch Output — Task queue with GPS tracking and evidence upload")

add_heading_section("6.9 Performance Evaluation")
add_body("Performance testing was conducted on the production deployment (Render.com free tier) with the following results:")

perf_data = [
    ("Homepage TTFB", "~0.5s", "Gunicorn sync worker"),
    ("Static Asset Cache", "31,536,000s", "Immutable, version-busted"),
    ("HTML Cache (repeat)", "60s", "Browser + ETag 304"),
    ("RSS Feed Edge Cache", "3,600s", "s-maxage for CDN"),
    ("Brotli Compression", "Level 4", "All text responses"),
    ("Gzip Compression", "Level 6", "Fallback for non-Brotli"),
    ("Preload Hints", "4 Link headers", "Early hints for critical resources"),
    ("First Contentful Paint", "~0.5s", "With preload hints"),
]
add_table(["Metric", "Value", "Notes"], perf_data, col_widths=[2.0, 1.8, 2.7])

add_heading_section("6.10 Security Evaluation")
add_body("The security header implementation was evaluated against the nine recommended HTTP security headers:")
security_data = [
    ("Content-Security-Policy", "Present", "Restricts resource loading sources"),
    ("Strict-Transport-Security", "Present", "Enforces HTTPS with includeSubDomains"),
    ("X-Content-Type-Options", "Present", "nosniff prevents MIME-type confusion"),
    ("X-Frame-Options", "Present", "DENY prevents clickjacking"),
    ("X-XSS-Protection", "Present", "1; mode=block for legacy browsers"),
    ("Referrer-Policy", "Present", "strict-origin-when-cross-origin"),
    ("Permissions-Policy", "Present", "Restricts browser features"),
    ("Cross-Origin-Opener-Policy", "Present", "same-origin isolation"),
    ("Cross-Origin-Embedder-Policy", "Present", "require-corp for COOP"),
]
add_table(["Header", "Status", "Purpose"], security_data, col_widths=[2.5, 1.0, 3.0])

add_heading_section("6.11 Accessibility Evaluation")
add_body("The application implements the following accessibility features:")
add_body("\u2022  Skip-to-content link for keyboard navigation")
add_body("\u2022  ARIA landmarks (navigation, main, contentinfo) on all pages")
add_body("\u2022  Semantic HTML5 elements (header, nav, main, footer)")
add_body("\u2022  Sufficient color contrast ratios (minimum 4.5:1 for normal text)")
add_body("\u2022  Form labels associated with inputs")
add_body("\u2022  Keyboard-navigable interactive elements")
add_body("\u2022  Alt text for images and icons")
add_body("\u2022  Screen reader announcements via aria-live regions")


# ════════════════════════════════════════════════════════════════════
# CHAPTER 7 — IMPACT ASSESSMENT
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("7. IMPACT ASSESSMENT")

add_body("The impact assessment evaluates the potential benefits of the SmartGarbage system across multiple dimensions. Note: As this is a prototype deployment, the impact values below are estimated projections based on system capabilities and comparable deployments, not measured results from a community-wide pilot.")

add_heading_section("7.1 Social Impact")
add_body("\u2022  Citizen Empowerment: The platform provides citizens with transparent access to collection schedules, complaint tracking, and ward-level performance metrics, enabling informed participation in waste governance.")
add_body("\u2022  Accessibility: Bilingual support (English and Telugu) ensures the system is accessible to residents with varying literacy levels. PWA offline support addresses connectivity limitations in semi-urban areas.")
add_body("\u2022  Accountability: Digital complaint tracking with status timelines and SLA monitoring creates accountability for complaint resolution.")

add_heading_section("7.2 Operational Impact")
add_body("Estimated operational improvements based on system capabilities:")
impact_data = [
    ("Schedule Accessibility", "100% digital access", "Previously verbal/informal"),
    ("Complaint Transparency", "Real-time tracking", "Previously untracked"),
    ("Worker Dispatch", "GPS-tracked assignment", "Previously manual coordination"),
    ("Data-Driven Decisions", "ML prediction + analytics", "Previously experience-based"),
]
add_table(["Metric", "With SmartGarbage", "Without SmartGarbage"], impact_data, col_widths=[2.0, 2.2, 2.3])

add_heading_section("7.3 Environmental Impact")
add_body("The environmental impact is estimated based on the system's capabilities:")
add_body("\u2022  Waste Segregation Incentive: Green Points gamification encourages source segregation, potentially increasing recycling rates.")
add_body("\u2022  Overflow Prevention: IoT monitoring and ML prediction enable proactive collection, reducing bin overflow and associated environmental contamination.")
add_body("\u2022  Carbon Footprint: Proper segregation reduces organic waste sent to landfills, potentially lowering methane emissions. Estimated CO2 savings are calculated conservatively at 0.5 kg CO2 per kg of recycled material.")

add_heading_section("7.4 Economic Feasibility")
add_body("\u2022  Zero-Cost Deployment: The system runs on Render.com's free tier with Supabase's free PostgreSQL, making it accessible to resource-constrained panchayats.")
add_body("\u2022  PAYT Revenue: Usage-based billing can generate revenue for the panchayat while incentivizing segregation.")
add_body("\u2022  Operational Cost Reduction: Optimized collection routes (enabled by IoT data) can reduce fuel and labor costs.")
add_body("\u2022  Open Source: The entire codebase is open source, eliminating vendor lock-in and enabling community contributions.")

add_heading_section("7.5 Scalability")
add_body("The architecture supports horizontal scaling through gunicorn workers and Redis-backed session/cache sharing. The system can be extended to additional panchayats by adding ward configurations to the WARD_COORDINATES mapping. Database schema supports multi-tenant isolation through ward-based filtering.")


# ════════════════════════════════════════════════════════════════════
# CHAPTER 8 — CHALLENGES FACED
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("8. CHALLENGES FACED")

challenges = [
    ("Data Unavailability", "Insufficient historical waste telemetry for ML training", "Generated synthetic dataset of 600 rows to demonstrate the prediction pipeline; documented the limitation and proposed real-data integration for future work"),
    ("Offline-First Architecture", "Implementing complaint filing without internet connectivity", "Implemented Service Worker with IndexedDB queue and Background Sync API for automatic submission on reconnection"),
    ("Real-Time Updates", "Live dashboard updates without page refresh", "Integrated Flask-SocketIO with Redis message broker for cross-worker broadcast of complaint and IoT updates"),
    ("IoT Integration", "Simulating sensor data without physical hardware", "Designed authenticated API endpoints for IoT telemetry ingestion; created test data generators for demonstration"),
    ("Security Hardening", "Achieving comprehensive security without compromising usability", "Implemented all nine OWASP-recommended headers; used Flask-Talisman for automatic header injection; balanced CSP policy to allow necessary resources"),
    ("Performance Optimization", "Reducing TTFB on free-tier hosting", "Implemented multi-layer caching (in-process TTL + Redis + browser ETag), Brotli compression, and preload hints for critical resources"),
    ("Bilingual Support", "Maintaining consistent translations across all templates", "Implemented Flask-Babel integration with gettext markers; created translation files for English and Telugu"),
    ("Push Notification Delivery", "Reliable delivery without a dedicated push service", "Implemented pywebpush with VAPID authentication; added subscription management and preference-based filtering"),
    ("Deployment Reliability", "Ensuring consistent behavior across development and production", "Implemented Docker containerization; used environment variable configuration; added comprehensive health check endpoint"),
    ("Payment Integration", "Testing Razorpay integration without live merchant account", "Implemented in test mode with Razorpay test keys; designed the flow to work with UPI as primary method"),
]

for challenge, problem, solution in challenges:
    add_heading_subsection(challenge)
    add_body(f"Problem: {problem}")
    add_body(f"Solution: {solution}")


# ════════════════════════════════════════════════════════════════════
# CHAPTER 9 — CONCLUSION
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("9. CONCLUSION")

add_heading_section("9.1 Summary")
add_body("This project presents SmartGarbage Chintalavalasa, a community-based smart waste management and digital governance system designed for the Chintalavalasa Gram Panchayat. The system integrates a Flask-based web application with PostgreSQL database, IoT smart-bin telemetry, machine learning prediction, PWA offline capabilities, PAYT billing, and comprehensive security into a unified platform.")

add_heading_section("9.2 Achievement of Objectives")
add_body("The project has successfully achieved the following objectives:")
add_body("\u2022  Developed a web-based platform providing waste collection schedules, complaint filing, and tracking accessible to all citizens.")
add_body("\u2022  Implemented an admin dashboard for real-time monitoring of complaints, worker dispatch, and analytics.")
add_body("\u2022  Designed IoT smart-bin telemetry integration for monitoring fill levels, battery status, and environmental conditions.")
add_body("\u2022  Developed a machine learning module using RandomForest regression for overflow risk prediction with transparent fallback.")
add_body("\u2022  Implemented PAYT billing with UPI/Razorpay integration and Green Points gamification.")
add_body("\u2022  Provided PWA capabilities with offline support, service worker, and install prompt.")
add_body("\u2022  Implemented comprehensive security measures including all nine OWASP-recommended headers.")
add_body("\u2022  Supported bilingual communication in English and Telugu.")

add_heading_section("9.3 Actual Contributions")
add_body("The key contributions of this project include:")
add_body("\u2022  An integrated, open-source waste management platform combining citizen engagement, administrative oversight, worker coordination, IoT monitoring, and ML prediction in a single system.")
add_body("\u2022  An offline-first PWA architecture enabling complaint filing and schedule access in low-connectivity areas.")
add_body("\u2022  A PAYT billing mechanism with gamification (Green Points) to incentivize waste segregation.")
add_body("\u2022  A comprehensive security implementation achieving all nine recommended HTTP security headers.")
add_body("\u2022  A deployable, low-cost architecture running entirely on free-tier services, suitable for resource-constrained panchayats.")

add_heading_section("9.4 Limitations of the Current Prototype")
add_body("The following limitations should be acknowledged:")
add_body("\u2022  The ML model is trained on synthetic data and has not been validated with real-world waste telemetry.")
add_body("\u2022  IoT telemetry is simulated; no physical smart bins are deployed.")
add_body("\u2022  The PAYT payment integration operates in test mode without a live merchant account.")
add_body("\u2022  Impact assessment values are estimated projections, not measured results from a community pilot.")
add_body("\u2022  The system has not been tested with real users at community scale.")
add_body("\u2022  SMS/WhatsApp notifications require third-party API keys not yet configured in production.")


# ════════════════════════════════════════════════════════════════════
# CHAPTER 10 — FUTURE WORK
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("10. FUTURE WORK")

add_body("The following improvements and extensions are planned for future development:")

add_heading_section("10.1 Real-World Data Integration")
add_body("\u2022  Replace synthetic ML training data with real historical waste telemetry from deployed IoT sensors.")
add_body("\u2022  Conduct a community pilot with actual residents to measure real-world impact metrics.")
add_body("\u2022  Collect baseline data before deployment to enable before-and-after comparison studies.")

add_heading_section("10.2 IoT Hardware Deployment")
add_body("\u2022  Deploy physical ultrasonic fill-level sensors on community bins.")
add_body("\u2022  Integrate GPS trackers on collection vehicles for route optimization.")
add_body("\u2022  Implement LoRaWAN or NB-IoT connectivity for low-power, wide-area sensor communication.")

add_heading_section("10.3 Mobile Application")
add_body("\u2022  Develop native Android application for broader accessibility.")
add_body("\u2022  Implement WhatsApp Business API integration for complaint filing via chatbot.")
add_body("\u2022  Add SMS-based schedule notifications for feature phone users.")

add_heading_section("10.4 Advanced Analytics")
add_body("\u2022  Implement time-series forecasting for waste generation patterns.")
add_body("\u2022  Add route optimization using Google OR-Tools or similar solvers.")
add_body("\u2022  Deploy computer vision for automated waste classification from citizen photos.")
add_body("\u2022  Create data visualization dashboards for panchayat administrators.")

add_heading_section("10.5 Multi-Panchayat Deployment")
add_body("\u2022  Extend the system to support multiple panchayats with tenant isolation.")
add_body("\u2022  Implement district-level aggregation and reporting.")
add_body("\u2022  Integrate with government APIs for SBM Grameen reporting compliance.")

add_heading_section("10.6 Payment and Billing")
add_body("\u2022  Activate live Razorpay merchant account for production payments.")
add_body("\u2022  Implement recurring billing and auto-debit via UPI mandates.")
add_body("\u2022  Add payment analytics and revenue reporting for panchayat finance.")


# ════════════════════════════════════════════════════════════════════
# REFERENCES
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("REFERENCES")
doc.add_paragraph()

references = [
    '[1] GOV.UK, "Design System," Government Digital Service, 2024. https://design-system.service.gov.uk/',
    '[2] Ministry of Housing and Urban Affairs, "Swachh Bharat Mission — Urban," Government of India, 2024. https://sbmurban.org/',
    '[3] Ministry of Jal Shakti, "Swachh Bharat Mission — Grameen Phase II," Government of India, 2024. https://sbmg.dop.gov.in/',
    '[4] T. Anagnostopoulos, A. Medvedev, A. Zaslavsky, and S. S. Kolomvatsos, "Waste Management as an IoT-Enabled Service in Smart Cities," in Internet of Things, Smart Spaces, and Next Generation Networks and Systems, Springer, 2015, pp. 104-115.',
    '[5] M. A. ad Din, A. R. Ghazali, and S. Hashim, "Smart Bin: IoT-Based Waste Monitoring System," in Proc. International Conference on Information Technology, 2020, pp. 1-6.',
    '[6] S. Kumar and R. Sharma, "IoT-Based Smart Waste Management System: A Review," Journal of Cleaner Production, vol. 295, 2021. doi:10.1016/j.jclepro.2021.126494',
    '[7] M. Afshin, B. Afshar, and A. Mohammadi, "Machine Learning Approaches for Municipal Solid Waste Generation Prediction," Waste Management, vol. 125, pp. 1-12, 2021.',
    '[8] Y. Chen, M. Chen, and H. Lin, "Predicting Municipal Solid Waste Generation Using Machine Learning Methods," Environmental Science and Pollution Research, vol. 27, 2020.',
    '[9] D. Thung and M. Yang, "Waste Classification using Convolutional Neural Network," in Proc. International Conference on Advanced Intelligent Systems and Informatics, Springer, 2016.',
    '[10] W3C, "Web Content Accessibility Guidelines (WCAG) 2.1," World Wide Web Consortium, 2018. https://www.w3.org/TR/WCAG21/',
    '[11] OWASP Foundation, "OWASP Top Ten 2021," Open Web Application Security Project, 2021. https://owasp.org/Top10/',
    '[12] U.S. Environmental Protection Agency, "Planning for a Sustainable Future: Pay-As-You-Throw," EPA, 2023. https://www.epa.gov/payt',
    '[13] Flask Documentation, "Flask — Pallets Projects," 2024. https://flask.palletsprojects.com/',
    '[14] SQLAlchemy Documentation, "SQLAlchemy — SQLAlchemy," 2024. https://docs.sqlalchemy.org/',
    '[15] scikit-learn Documentation, "scikit-learn: Machine Learning in Python," 2024. https://scikit-learn.org/',
    '[16] MDN Web Docs, "Progressive Web Apps (PWAs)," Mozilla Developer Network, 2024. https://developer.mozilla.org/en-US/docs/Web/Progressive_web_apps',
    '[17] Web Push Protocol, "Push API — MDN Web Docs," Mozilla Developer Network, 2024. https://developer.mozilla.org/en-US/docs/Web/API/Push_API',
    '[18] Supabase Documentation, "Supabase — The Open Source Firebase Alternative," 2024. https://supabase.com/docs',
    '[19] Render Documentation, "Render — Cloud Application Platform," 2024. https://render.com/docs',
    '[20] Razorpay Documentation, "Razorpay Docs," 2024. https://docs.razorpay.com/',
]
for ref in references:
    add_paragraph(ref, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)


# ════════════════════════════════════════════════════════════════════
# APPENDIX A — PACKAGES, TOOLS USED & WORKING PROCESS
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("APPENDIX A: PACKAGES, TOOLS USED & WORKING PROCESS")

add_heading_section("A.1 Packages and Tools Used")
packages = [
    ("Flask 3.1.3", "Python", "Web framework"),
    ("Flask-SQLAlchemy 3.1.1", "Python", "ORM integration"),
    ("Flask-Login 0.6.3", "Python", "User session management"),
    ("Flask-Talisman 1.1.0", "Python", "Security headers"),
    ("Flask-Limiter 4.1.1", "Python", "Rate limiting"),
    ("Flask-Compress 1.17", "Python", "Brotli/Gzip compression"),
    ("Flask-WTF 1.2.1", "Python", "Form handling and CSRF"),
    ("Flask-Mailman 1.1.1", "Python", "Email notifications"),
    ("Flask-SocketIO 5.3.6", "Python", "Real-time WebSocket communication"),
    ("Flask-Session 0.8.0", "Python", "Server-side sessions"),
    ("Flask-Migrate 3.1.0", "Python", "Database migrations"),
    ("SQLAlchemy 2.0.50", "Python", "Database ORM"),
    ("psycopg2-binary 2.9.10", "Python", "PostgreSQL adapter"),
    ("Redis 6.2.0", "Python", "Redis client"),
    ("RQ 2.2.0", "Python", "Background job queue"),
    ("scikit-learn 1.9.0", "Python", "Machine learning"),
    ("pandas 3.0.3", "Python", "Data manipulation"),
    ("numpy 2.4.6", "Python", "Numerical computing"),
    ("matplotlib 3.10.9", "Python", "Data visualization"),
    ("pywebpush 2.0.0", "Python", "Web push notifications"),
    ("reportlab 5.0.0", "Python", "PDF invoice generation"),
    ("sentry-sdk 2.5.0", "Python", "Error tracking"),
    ("gunicorn 26.0.0", "Python", "Production WSGI server"),
    ("Jinja2 3.1.6", "Python", "Template engine"),
    ("Bootstrap 5", "CSS/JS", "UI framework"),
    ("Socket.IO", "JavaScript", "Real-time client communication"),
    ("Docker", "DevOps", "Containerization"),
    ("Git", "DevOps", "Version control"),
    ("PostgreSQL (Supabase)", "Database", "Relational database"),
    ("Redis", "Database", "In-memory data store"),
    ("Render.com", "Cloud", "Application hosting"),
    ("Sentry", "Monitoring", "Error tracking"),
]
add_table(["Package/Tool", "Category", "Purpose"], packages, col_widths=[2.5, 1.2, 2.8])

add_heading_section("A.2 Working Process")
add_body("The development process followed these steps:")
add_body("1. Environment Setup: Installed Python 3.12, created a virtual environment, installed Flask and dependencies. Set up PostgreSQL database on Supabase and Redis on Upstash.")
add_body("2. Project Structure: Organized the codebase into modular directories — app/routes/ for route modules, app/templates/ for Jinja2 templates, app/static/ for CSS/JS/images, and app/models.py for SQLAlchemy models.")
add_body("3. Version Control: Used Git with GitHub for version control. The project repository is maintained at github.com/jaganmohan08112005-sketch/SmartgarbageCSP.")
add_body("4. Iterative Development: Developed each module incrementally — public portal first, then citizen portal, admin portal, worker portal, IoT integration, ML prediction, and finally PWA/offline capabilities.")
add_body("5. Testing: Conducted functional testing of all routes and features using Flask's test client. Automated syntax verification using python -m compileall. Security header verification using curl.")
add_body("6. Deployment: Containerized the application using Docker with a multi-stage build. Deployed to Render.com with automatic builds from GitHub main branch pushes.")
add_body("7. Monitoring: Integrated Sentry for production error tracking. Implemented a health check endpoint (/health) for uptime monitoring.")


# ════════════════════════════════════════════════════════════════════
# APPENDIX B — IMPORTANT SOURCE CODE
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("APPENDIX B: IMPORTANT SOURCE CODE")

add_heading_section("B.1 Application Factory (app/__init__.py)")
add_body("The Flask application factory initializes all extensions, configures security headers, sets up caching, and registers background jobs. Key components include Flask-Talisman for security headers, Flask-Limiter for rate limiting, Flask-SocketIO for real-time updates, and multi-layer caching with in-process TTL and Redis.")

add_heading_section("B.2 Database Models (app/models.py)")
add_body("The database models define the system's data entities: User (with role-based authentication), Complaint (with GPS and photo support), SmartBin (IoT telemetry), WorkerProfile (GPS tracking), WasteDeclaration (segregation reporting), PAYTInvoice (billing), PushSubscription and NotificationPreference (push notifications), and ConsentRecord (GDPR compliance).")

add_heading_section("B.3 ML Prediction Module (app/ml_model.py)")
add_body("The ML module implements a RandomForest regressor trained on synthetic data to predict bin overflow risk. Features include day-of-week, season index, recent complaint volume, and ward identifier. The module includes a transparent heuristic fallback when the model artifact is unavailable, ensuring the route never errors.")

add_heading_section("B.4 Service Worker (app/static/sw.js)")
add_body("The service worker implements a versioned precache manifest for offline support. It uses a stale-while-revalidate strategy for HTML pages and cache-first for static assets. The push event handler processes web push notifications and displays them to the user with action buttons.")

add_heading_section("B.5 Push Notification Module (app/push.py)")
add_body("The push notification module manages VAPID key generation, subscription storage, preference-based filtering, and delivery logging. It integrates with the complaint lifecycle to automatically notify citizens when their complaint status changes, respecting their notification preferences.")


# ════════════════════════════════════════════════════════════════════
# PAPER PUBLICATIONS
# ════════════════════════════════════════════════════════════════════
add_page_break()
add_heading_chapter("PAPER PUBLICATIONS")
doc.add_paragraph()
add_body("No research papers have been published based on this project at the time of report submission. The project results and findings may be submitted for publication in future.")


# ════════════════════════════════════════════════════════════════════
# SAVE
# ════════════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "SmartGarbage_Community_Project_Report.docx")
doc.save(output_path)
print(f"Report generated successfully: {output_path}")
print(f"File size: {os.path.getsize(output_path):,} bytes ({os.path.getsize(output_path) / 1024:.1f} KB)")
